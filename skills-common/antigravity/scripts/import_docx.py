#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
import_docx.py — .docx  ->  .md, va phat hien file bi sua tay.

Ba viec:
  1. Di cu file .docx cu sang .md
  2. Cuu ho khi co nguoi sua truc tiep trong Word
  3. Che do soat: nhan .docx bat ky, dua ve .md de validate.py doc duoc

Day KHONG phai bai toan doc Word tu do. De cuong da khoa hinh dang moi bang, nen
doc theo khuon da biet: ten muc biet truoc, so cot biet truoc. Cho nao khong khop
khuon thi BAO LOI, khong doan.

    python import_docx.py cu.docx -o FUNC-QLNSD-001.md
    python import_docx.py sua-tay.docx --diff FUNC-QLNSD-001.md
"""
from __future__ import annotations

import argparse
import difflib
import hashlib
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import srslib as S


# ---------------------------------------------------------------------------
def heading_level(p) -> int | None:
    m = re.fullmatch(r"Heading (\d)", p.style.name or "")
    return int(m.group(1)) if m else None


def iter_body(doc: Document):
    """Paragraphs and tables in document order — python-docx exposes them
    separately, so walk the XML instead."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def inline_md(paragraph) -> str:
    """Rebuild inline markdown from a paragraph's runs.

    render.py turns `**text**` into bold runs, so import has to put the markers
    back or the round trip loses the emphasis and the two sources drift.
    """
    out = []
    for r in paragraph.runs:
        t = r.text
        if not t:
            continue
        if r.bold and t.strip():
            lead = t[:len(t) - len(t.lstrip())]
            trail = t[len(t.rstrip()):]
            out.append(f"{lead}**{t.strip()}**{trail}")
        else:
            out.append(t)
    joined = "".join(out).strip()
    # Adjacent bold runs would come back as `**a****b**`; collapse them.
    return re.sub(r"\*\*(\s*)\*\*", r"\1", joined)


BULLET_SEP = "·"
_ENDS_SENTENCE = ".:;!?…,"


def _bullet_level(par, bullet_styles: dict[str, int] | None) -> int:
    """Bullet depth of a paragraph: 0 = not a bullet, 1–3 = level.

    Two mechanisms, because real documents use both. Word's own numbering
    (`w:numPr`) is the obvious one. But specs written to this organisation's
    template carry no numbering at all — they use the house bullet styles
    `T-Gach -` / `T-Gach +` / `T-Gach *`, which the outline already names as
    `bullet_1..3`. Checking only `w:numPr` therefore missed every bullet in an
    actual legacy spec while looking correct on synthetic ones.

    Generic style names are deliberately not consulted: `List Paragraph` is
    routinely applied to plain indented text, and trusting it would invent
    bullets that were never there.
    """
    if bullet_styles:
        lvl = bullet_styles.get(par.style.name or "")
        if lvl:
            return lvl
    pPr = par._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        ilvl = pPr.find(qn("w:numPr")).find(qn("w:ilvl"))
        try:
            return min(3, int(ilvl.get(qn("w:val"))) + 1) if ilvl is not None else 1
        except (TypeError, ValueError):
            return 1
    return 0


def bullet_style_map(outline: dict) -> dict[str, int]:
    """`{style name: level}` for the outline's three bullet styles."""
    return {outline["styles"][f"bullet_{n}"]: n
            for n in (1, 2, 3) if outline["styles"].get(f"bullet_{n}")}


def cell_text(cell, bullet_styles: dict[str, int] | None = None) -> str:
    """One markdown-safe line, with the cell's internal structure still legible.

    A markdown table row cannot span lines, so a multi-paragraph Word cell has
    to collapse. Collapsing with a bare space — what this did originally — runs
    unpunctuated list items together into an unreadable wall: five input rules
    became one 300-character sentence, and the reviewer could no longer see
    that there had been five of them to split into five `BR-` codes.

    So the boundaries are kept as marks instead of dropped:
      · bulleted item  → prefixed with `·`
      · ordinary line  → given a full stop when it lacks end punctuation

    Two invariants, both load-bearing, both covered by evals: the result never
    contains a newline (a soft line break — Shift+Enter — used to survive here
    and split the markdown row in half, silently truncating the table), and
    never a bare `|` (which would split one cell into several).
    """
    segs: list[tuple[str, int]] = []
    for p in cell.paragraphs:
        lvl = _bullet_level(p, bullet_styles)
        # `python-docx` renders <w:br/> as \n inside a single paragraph, so a
        # paragraph can itself hold several visual lines.
        for piece in p.text.split("\n"):
            piece = piece.strip()
            if piece:
                segs.append((piece, lvl))
    if not segs:
        return ""

    out: list[str] = []
    for i, (txt, lvl) in enumerate(segs):
        if lvl:
            # Depth shown by repeating the mark, so a nested list stays
            # distinguishable from a flat one after collapsing.
            out.append(f"{BULLET_SEP * lvl} {txt}")
        elif i < len(segs) - 1 and txt[-1] not in _ENDS_SENTENCE:
            out.append(txt + ".")
        else:
            out.append(txt)
    line = " ".join(out)
    # Escape after joining: a literal pipe anywhere would end the cell early.
    return line.replace("|", r"\|")


def cell_segment_count(cell) -> int:
    """How many visual lines the cell held before being collapsed."""
    n = 0
    for p in cell.paragraphs:
        n += sum(1 for piece in p.text.split("\n") if piece.strip())
    return n


def images_in(paragraph) -> list[tuple[str, str]]:
    """(rel_id, stored_name) for every inline picture in the paragraph."""
    out = []
    for docPr in paragraph._p.iter(qn("wp:docPr")):
        name = docPr.get("descr") or docPr.get("name") or ""
        blip = None
        parent = docPr.getparent().getparent()
        for b in parent.iter(qn("a:blip")):
            blip = b.get(qn("r:embed"))
            break
        if blip:
            out.append((blip, name))
    return out


# ---------------------------------------------------------------------------
class Importer:
    def __init__(self, outline: dict, assets: Path, strict: bool = True):
        self.o = outline
        self.lex = outline["lexicon"]
        self.assets = assets
        self.strict = strict
        self.problems: list[str] = []
        self.saved: list[str] = []
        # Cells that held several lines in Word. They survive as one line with
        # `·` marks, but the count is worth reporting: in a constraint column,
        # five bullets in one cell is usually five business rules that should
        # each get their own `BR-` row, and that is a judgement only the BA can
        # make. Silently collapsing them hid the decision entirely.
        self.merged_cells: list[tuple[str, int, str]] = []
        # rel_id -> where the picture sits, so a legacy image can be named
        # after its section instead of Word's meaningless `Picture 1`.
        self._img_index: dict = {}
        self._img_fallback = 0
        self.is_group = False
        self.group_names: set[str] = set()   # nhóm không còn mục con

    def run(self, src: Path) -> tuple[dict, list[str]]:
        doc = Document(str(src))
        base = self._detect_levels(doc)
        meta = self._meta(doc, src)
        self._img_index = {h.rel_id: h
                           for h in S.scan_images(doc, self.o["styles"])}

        lines: list[str] = []
        table_i = 0
        pending_caption: str | None = None
        cur_section = ""
        auto_names = self._auto_sections()

        for item in iter_body(doc):
            if hasattr(item, "rows"):
                # Only a fully generated section goes back to front matter.
                # `column:` and `registry:` mark sections where a script fills
                # one column but the BA writes the rest — those tables stay.
                if auto_names.get(cur_section) == "changelog":
                    meta["changelog"] = self._changelog(item)
                    table_i += 1
                    continue
                lines += self._table(item)
                table_i += 1
                continue

            lvl = heading_level(item)
            text = item.text.strip()

            if lvl is not None and text:
                mg = re.match(r"Nhóm chức năng\s*\[([^\]]+)\]\s*(.*)", text)
                if mg:
                    self.is_group = True
                    cur_section = S.GROUP_DESC
                    meta["ma"], meta["ten"] = mg.group(1).strip(), \
                        mg.group(2).strip()
                    meta["profile"] = "GROUP"
                    lines += [f"# {text}", ""]
                    cur_section = ""
                    continue
                if lvl == base:
                    m = re.match(r"Chức năng\s*\[([^\]]+)\]\s*(.*)", text)
                    if m:
                        meta["ma"], meta["ten"] = m.group(1).strip(), \
                            m.group(2).strip()
                    lines += [f"# {text}", ""]
                    cur_section = ""
                elif lvl == base + 1:
                    lines += [f"## {text}", ""]
                    cur_section = text
                elif lvl == base + 2:
                    lines += [f"### {text}", ""]
                    cur_section = text
                else:
                    self.problems.append(
                        f"Heading {lvl} “{text[:40]}” nằm ngoài ba cấp đề cương "
                        f"({base}/{base+1}/{base+2}).")
                continue

            imgs = images_in(item)
            if imgs:
                for rid, name in imgs:
                    path = self._save_image(doc, rid, name)
                    lines += [f"![{pending_caption or ''}]({path})", ""]
                pending_caption = None
                continue

            if not text:
                continue

            style = item.style.name or ""
            if style == self.o["styles"]["caption"]:
                # A caption follows its picture, so attach it to the line just
                # emitted rather than waiting for the next one.
                self._attach_caption(lines, text)
                continue
            if style == self.o["styles"]["note"]:
                # In a function file T-GhiChu is guidance for the BA and is
                # dropped. In a group file it is the description itself —
                # dropping it there loses the only content the file has.
                if self.is_group:
                    lines += [inline_md(item), ""]
                continue
            bullets = {self.o["styles"].get(f"bullet_{n}"): n
                       for n in (1, 2, 3) if self.o["styles"].get(f"bullet_{n}")}
            if style in bullets:
                lines += ["  " * (bullets[style] - 1) + f"- {text}"]
                continue
            if self.lex["missing_image"] in text:
                continue                      # placeholder box, not content

            runs = [r for r in item.runs if r.text.strip()]
            if runs and all(r.bold for r in runs):
                # In a group document, bold body IS the section label — the
                # group uses Heading 2 and its children cannot use Heading 3,
                # which the function titles already occupy.
                if self.is_group and S.norm(text) in self.group_names:
                    lines += [f"## {text}", ""]
                    cur_section = text
                    continue
                # Otherwise it is a table label; without the markers it would
                # come back as prose and the round trip would not close.
                lines += [f"**{text}**", ""]
                continue

            lines += [inline_md(item), ""]

        return meta, lines

    def _auto_sections(self) -> dict[str, str]:
        out = {}
        for prof in self.o["profiles"].values():
            for s in prof["function_sections"] + prof["feature_sections"]:
                if s.get("auto"):
                    out[s["name"]] = s["auto"]
        return out

    def _changelog(self, t) -> list[dict]:
        bs = bullet_style_map(self.o)
        rows = [[cell_text(c, bs) for c in r.cells] for r in t.rows]
        out = []
        for r in rows[1:]:
            if len(r) >= 4 and " ".join(r).strip():
                out.append({"v": r[0], "ngay": r[1], "nguoi": r[2],
                            "mo_ta": r[3]})
        return out

    # -- helpers -----------------------------------------------------------
    def _detect_levels(self, doc: Document) -> int:
        for p in doc.paragraphs:
            if heading_level(p) and p.text.strip().startswith("Nhóm chức năng"):
                return heading_level(p) + 1   # functions sit one level below

        lvls = sorted({heading_level(p) for p in doc.paragraphs
                       if heading_level(p) and p.text.strip()})
        if not lvls:
            self.problems.append("Không tìm thấy heading nào — file có đúng là "
                                 "tài liệu đặc tả chức năng không?")
            return 3
        return lvls[0]

    def _meta(self, doc: Document, src: Path) -> dict:
        cp = doc.core_properties
        meta = {"ma": "", "ten": "", "profile": "", "version": "0.1",
                "status": "draft"}
        note = cp.comments or ""
        m = re.search(r"md_sha=([0-9a-f]+)", note)
        if m:
            meta["_md_sha"] = m.group(1)
        m = re.search(r"doc_sha=([0-9a-f]+)", note)
        if m:
            meta["_doc_sha"] = m.group(1)

        texts = [p_.text for p_ in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                texts += [c.text for c in row.cells]
        meta["_doc_sha_now"] = S.content_sha(texts)
        m = re.search(r"version=([^\s|]+)", note)
        if m:
            meta["version"] = m.group(1)
        # Files rendered before the skill was renamed carry the old names.
        # "ntda-srs-helper" happens to contain "srs-help", but relying on that
        # coincidence would break the moment either name changes again.
        meta["_from_skill"] = any(
            n in note for n in ("srs-help", "ntda-srs-helper", "srs-writer"))
        return meta

    def _row_cells(self, row, bs) -> list[str]:
        """Cell texts for one row, with merged spans written once.

        `python-docx` resolves a horizontally merged cell by handing the same
        text back for every column it covers, so a divider band merged across
        six columns imported as `| Button | Button | Button | Button | Button |
        Button |` — and rendered back out exactly that way, six times, in a
        delivered document. Keep the text in the first column and blank the
        rest: the row still reads as a band, and nothing is duplicated.
        """
        out: list[str] = []
        seen: set[int] = set()
        for c in row.cells:
            key = id(c._tc)
            if key in seen:
                out.append("")
            else:
                seen.add(key)
                out.append(cell_text(c, bs))
        return out

    def _table(self, t) -> list[str]:
        bs = bullet_style_map(self.o)
        rows = [self._row_cells(r, bs) for r in t.rows]
        if not rows:
            return []
        # Seed the ordinals so the .md reads sensibly on its own. Render
        # recomputes them anyway, so a wrong number here cannot reach the
        # document — which is what makes seeding safe rather than a second
        # source of truth.
        rows = S.renumber_stt(rows)
        head = rows[0] if rows else []
        for ri, r in enumerate(t.rows):
            if ri == 0:
                continue
            for ci, c in enumerate(r.cells):
                n = cell_segment_count(c)
                if n > 1:
                    col = head[ci] if ci < len(head) else f"cột {ci + 1}"
                    self.merged_cells.append((col, n, c.text.split("\n")[0][:52]))
        # The missing-image box is rendered as a table; it is scaffolding, not
        # document content, so drop it rather than importing it back.
        if rows[0] and self.lex["missing_image"] in rows[0][0]:
            return []
        out = ["| " + " | ".join(rows[0]) + " |",
               "|" + "|".join(["---"] * len(rows[0])) + "|"]
        for r in rows[1:]:
            out.append("| " + " | ".join(c.replace("\n", " ") for c in r) + " |")
        out.append("")
        return out

    @staticmethod
    def _attach_caption(lines: list[str], text: str):
        text = re.sub(r"^(Hình|Bảng)\s*\d*\s*[.:]?\s*", "", text).strip()
        for i in range(len(lines) - 1, -1, -1):
            if lines[i].startswith("!["):
                lines[i] = re.sub(r"^!\[[^\]]*\]", f"![{text}]", lines[i])
                return
        if text:
            lines += [f"**{text}**", ""]

    # Word's default alt-text is `Picture 1` for nearly every image, so it is
    # only a filename when the skill itself put one there. Anything without a
    # picture extension is Word's own label and must not become a filename:
    # naming by it collapsed nineteen images onto two files, silently, while
    # still reporting nineteen.
    _ALT_IS_FILENAME = re.compile(r"\.(png|jpe?g|gif|bmp|tiff?|svg)$", re.I)

    def _save_image(self, doc: Document, rid: str, name: str) -> str:
        part = doc.part.related_parts.get(rid)
        if part is None:
            self.problems.append(f"Không lấy được ảnh (rel {rid}).")
            return "assets/KHONG-LAY-DUOC.png"

        if name and "/" not in name and self._ALT_IS_FILENAME.search(name):
            # Rendered by this skill: the alt-text is the original filename and
            # keeping it is what makes .md → .docx → .md round-trip clean.
            fname = name
        else:
            hit = self._img_index.get(rid)
            ext = S.IMAGE_EXT.get(getattr(part, "content_type", "")) \
                or Path(str(part.partname)).suffix or ".bin"
            if hit is not None:
                fname = S.image_name(hit.seq, hit.sec_num, hit.caption,
                                     hit.context, ext)
            else:
                self._img_fallback += 1
                fname = f"{self._img_fallback:03d}_khong-ro-vi-tri{ext}"

        self.assets.mkdir(parents=True, exist_ok=True)
        dst = S.unique_path(self.assets, fname, part.blob)
        dst.write_bytes(part.blob)
        self.saved.append(dst.name)
        return f"{self.assets.name}/{dst.name}"


# ---------------------------------------------------------------------------
def build_md(meta: dict, lines: list[str], outline: dict,
             profile: str | None) -> str:
    fm = {
        "ma": meta.get("ma", ""),
        "ten": meta.get("ten", ""),
        "profile": profile or meta.get("profile") or "",
        "version": meta.get("version", "0.1"),
        "status": meta.get("status", "draft"),
        "outline_id": outline["id"],
        "outline_version": outline["version"],
        "changelog": meta.get("changelog") or [
            {"v": meta.get("version", "0.1"), "ngay": "", "nguoi": "",
             "mo_ta": "Nhập từ .docx"}],
    }
    return S.dump_front_matter(fm) + "\n\n" + "\n".join(lines).rstrip() + "\n"


def guess_profile(lines: list[str], outline: dict) -> tuple[str | None, str]:
    """Match the headings against each profile and take the best fit.

    Function-level sections alone are not enough: DANHMUC has exactly the same
    ones as UI, and only differs at feature level. Score both tiers, and prefer
    the more specific profile when they tie.

    Returns (kind, report). A weak match returns None rather than the
    least-bad guess — a document that does not follow the outline needs to be
    told so, not quietly labelled.
    """
    fn = {l[3:].strip() for l in lines if l.startswith("## ")
          and not l.startswith("## Tính năng")}
    ft = {l[4:].strip() for l in lines if l.startswith("### ")}

    # A merged master document contains many child files at once. Importing it
    # as if it were a single child silently keeps the first heading and drops
    # everything else, which looks like a successful import.
    n_grp = sum(1 for l in lines if l.startswith("# Nhóm chức năng")
                or l.startswith("## Nhóm chức năng"))
    n_fn = sum(1 for l in lines if l.startswith("# Chức năng")
               or l.startswith("## Chức năng"))
    if n_grp + n_fn > 1:
        return None, (f"chứa {n_grp} nhóm và {n_fn} chức năng — đây là TÀI LIỆU "
                      f"TỔNG đã ghép, không phải một file con")

    # A group document is unmistakable: its title line names it. It has no
    # sections at all now — a heading plus an optional description — so there
    # is nothing else to match on.
    if any(l.startswith("# Nhóm chức năng") for l in lines):
        return S.GROUP, "tiêu đề là nhóm chức năng"

    ranked = []
    for name, prof in outline["profiles"].items():
        w_fn = {s["name"] for s in prof["function_sections"]}
        w_ft = {s["name"] for s in prof["feature_sections"]}
        if prof["has_features"] and not ft:
            continue
        if not prof["has_features"] and ft:
            continue
        hit = len(w_fn & fn) + len(w_ft & ft)
        total = len(w_fn) + len(w_ft)
        miss = len(w_fn - fn) + len(w_ft - ft)
        extra = len(fn - w_fn) + len(ft - w_ft)
        ranked.append((miss + extra, -hit, total, name, hit))

    if not ranked:
        return None, "Không có loại nào phù hợp về cấu trúc."
    ranked.sort()
    _, _, total, name, hit = ranked[0]
    rate = hit / total if total else 0
    report = f"khớp {hit}/{total} tên mục của loại {name}"

    if rate < 0.5:
        if hit == 0:
            return None, "không có tên mục nào trùng đề cương"
        return None, f"chỉ {report}, quá thấp để nhận dạng"
    return name, report


def main() -> int:
    S.utf8_stdio()
    ap = argparse.ArgumentParser(description="Nhập .docx về .md.")
    ap.add_argument("src")
    ap.add_argument("-o", "--out", default=None)
    ap.add_argument("--assets", default="assets")
    ap.add_argument("--profile", default=None)
    ap.add_argument("--diff", default=None,
                    help="so với file .md hiện có thay vì ghi ra file mới")
    ap.add_argument("--raw", action="store_true",
                    help="vẫn xuất .md dù file không theo đề cương; kết quả "
                         "không hợp lệ, chỉ để đọc và đối chiếu")
    ap.add_argument("--outline", default=None)
    a = ap.parse_args()

    outline = S.load_outline(a.outline)
    src = Path(a.src)
    if not src.exists():
        print(f"LỖI: không thấy {src}", file=sys.stderr)
        return 1

    out = Path(a.out) if a.out else src.with_suffix(".md")
    imp = Importer(outline, Path(a.assets) if a.diff is None
                   else Path(a.diff).parent / a.assets)
    meta, lines = imp.run(src)

    guessed, report = guess_profile(lines, outline)

    # This gate runs before --profile is even consulted. A merged master
    # document contains many child files at once; importing it as a single
    # child silently keeps the first heading and drops the rest, and that
    # looks exactly like a successful import. Passing --profile explicitly
    # used to skip straight past this check (profile was no longer None), so
    # the one flag meant to force a *type* could also be used, unintentionally
    # or not, to force past a structural gate it has nothing to do with. The
    # count of groups/functions is a property of the file, not of what the
    # caller believes it is — so it is checked unconditionally here.
    if "TÀI LIỆU TỔNG" in report:
        print(f"\nĐÂY LÀ TÀI LIỆU TỔNG — {report}.\n")
        print("Skill làm việc trên từng file con, không tách tài liệu tổng.")
        print("Dùng `tools/split_master.py` của pipeline để tách ra trước, rồi "
              "nhập từng file.")
        print("Cổng này áp dụng bất kể có truyền --profile hay không — file "
              "gộp nhiều nhóm/chức năng không nhập được như một file con.")
        return 1

    profile = a.profile or guessed
    if profile is None and a.raw:
        print(f"\nFile không theo đề cương ({report}).")
        print("--raw: vẫn xuất bản thô để đọc đối chiếu. Kết quả KHÔNG hợp lệ, "
              "đừng sửa dần\nthành bản chuẩn — xem mục “Chuyển tài liệu cũ”.")
        profile = a.profile or "UI"
    elif profile is None:
        print(f"\nFILE KHÔNG THEO ĐỀ CƯƠNG HIỆN HÀNH — {report}\n")
        print("Đây là tài liệu cũ hoặc viết theo mẫu khác. Không nhập tự động "
              "được vì\nkhông biết nội dung nào thuộc mục nào của khung mới.")
        print("\nCách làm: xem mục “Chuyển tài liệu cũ” trong SKILL.md — dựng "
              "khung mới bằng\nscaffold.py rồi chuyển nội dung sang từng mục, "
              "có Claude soát từng bước.")
        print("\nNếu bạn vẫn muốn lấy phần thô ra để đối chiếu, chạy lại kèm "
              "--profile «loại»\nvà --raw. Kết quả sẽ KHÔNG hợp lệ, chỉ để "
              "đọc.")
        return 1

    md = build_md(meta, lines, outline, profile)

    # -- rescue mode ------------------------------------------------------
    if a.diff:
        ref = Path(a.diff)
        cur = ref.read_text(encoding="utf-8")
        cur_sha = hashlib.sha256(ref.read_bytes()).hexdigest()[:16]
        md_sha = meta.get("_md_sha")
        doc_sha, doc_now = meta.get("_doc_sha"), meta.get("_doc_sha_now")

        print(f"Nguồn .docx : {src}")
        print(f"Bản .md hiện: {ref}\n")

        edited = doc_sha is not None and doc_sha != doc_now
        md_moved = md_sha is not None and md_sha != cur_sha

        if not meta.get("_from_skill"):
            print("→ File .docx này KHÔNG do skill render. Coi như file ngoài "
                  "luồng — nhập về rồi tự đối chiếu.")
        elif doc_sha is None:
            print("→ File render bằng bản skill cũ, chưa có dấu vết nội dung. "
                  "Không kết luận được có bị sửa tay hay không.")
        elif edited and md_moved:
            print("→ CẢ HAI ĐỀU ĐỔI. Có người sửa trong Word, VÀ bản .md cũng "
                  "đã thay đổi sau lần render đó. Phải nhập thủ công từng "
                  "điểm.")
        elif edited:
            print("→ FILE .DOCX BỊ SỬA TAY. Nội dung trong Word khác với thứ "
                  "skill đã render.")
        elif md_moved:
            print("→ .docx sạch, nhưng bản .md đã thay đổi sau lần render. "
                  "Chỉ cần render lại, không cần nhập.")
        else:
            print("→ Sạch: .docx đúng bản render của .md hiện tại, không ai "
                  "sửa tay.")

        # Compare bodies only. Front matter differs by construction — import
        # cannot recover `nhom` or `status` from a .docx — and letting that
        # noise into the diff would bury the real edits.
        _, cur_body = S.parse_front_matter(cur)
        _, new_body = S.parse_front_matter(md)
        # Blank lines carry no meaning in this format; keeping them would
        # report cosmetic noise as edits.
        strip = lambda t: [l for l in t.splitlines() if l.strip()]
        d = list(difflib.unified_diff(
            strip(cur_body), strip(new_body),
            fromfile=f"{ref} (thân bài)", tofile=f"{src} (nhập về)",
            lineterm="", n=1))
        body = [l for l in d if l.startswith(("+", "-"))
                and not l.startswith(("+++", "---"))]

        if not body:
            print("\nThân bài trùng khớp hoàn toàn.")
        else:
            print(f"\n{len(body)} dòng khác biệt ở thân bài:\n")
            print("\n".join(d[:120]))
            if len(d) > 120:
                print(f"… còn {len(d)-120} dòng nữa")

        print("\nKHÔNG ghi đè gì cả. Bạn quyết định giữ bên nào rồi sửa .md "
              "bằng tay.")
        print("Lưu ý: front matter không nằm trong phép so — .docx không mang "
              "đủ thông tin để dựng lại.")
    else:
        out.write_text(md, encoding="utf-8")
        print(f"OK -> {out}   (loại {profile}"
              + (", đoán từ tên mục)" if not a.profile else ")"))
        if not a.profile:
            print(f"  {report}")

    if imp.saved:
        print(f"  Ảnh lấy ra: {len(imp.saved)} tệp -> {imp.assets}/")
        for n in imp.saved:
            print(f"    {n}")
        # Count what the package holds against what reached disk. This is the
        # check that was missing when nineteen images silently became two
        # files: the run reported nineteen either way.
        try:
            import zipfile
            import re as _re
            with zipfile.ZipFile(src) as z:
                rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
                body = z.read("word/document.xml").decode("utf-8")
            # Only pictures the body actually references. `word/media/` also
            # holds whatever the template carries — base.docx ships an unused
            # image — and counting those reported losses that never happened.
            rmap = dict(_re.findall(r'Id="([^"]+)"[^>]*Target="([^"]+)"', rels))
            used = {rmap.get(r, "") for r in _re.findall(r'r:embed="([^"]+)"', body)}
            media = [m for m in used if m]
            vec = [m for m in media
                   if Path(m).suffix.lower() in S.VECTOR_EXT]
            raster = len(media) - len(vec)
            on_disk = len({n for n in imp.saved})
            if on_disk < raster:
                print(f"    [LỖI ] gói .docx chứa {raster} ảnh bitmap nhưng "
                      f"chỉ ghi ra {on_disk} tệp — có ảnh bị mất.")
            if vec:
                print(f"    [CẢNH] {len(vec)} ảnh định dạng vector "
                      f"({', '.join(sorted({Path(m).suffix for m in vec}))}) "
                      f"chưa lấy được — mở Word, chuột phải ảnh → Save as "
                      f"Picture, rồi đặt vào {imp.assets}/.")
        except Exception:
            pass
    for p in imp.problems:
        print(f"  CẢNH BÁO: {p}")

    if imp.merged_cells:
        # Loud on purpose. A row in Word held several lines; markdown tables
        # cannot, so they are now one line separated by `·`. Nothing is lost,
        # but a constraint cell holding five rules is usually five business
        # rules that each want their own `BR-` row — a call the BA has to make
        # and could not make while the collapse was silent.
        from collections import Counter
        by_col = Counter(col for col, _, _ in imp.merged_cells)
        worst = sorted(imp.merged_cells, key=lambda x: -x[1])[:5]
        print(f"\n  CẢNH BÁO: {len(imp.merged_cells)} ô trong Word có nhiều "
              f"dòng, đã gộp thành một dòng ngăn bằng `·`"
              f" (bảng markdown không xuống dòng trong ô được).")
        for col, n in by_col.most_common(5):
            print(f"    cột “{col}”: {n} ô")
        print("    nhiều dòng nhất:")
        for col, n, first in worst:
            print(f"      {n} dòng · {col} · {first}…")
        print("    → Soát lại: ô mô tả ràng buộc gộp nhiều ý thường nên tách "
              "thành các mã BR- riêng, mỗi quy tắc một mã.")

    if not a.diff:
        print("\n  Bước tiếp: chạy validate.py. Front matter mới nhập còn "
              "thiếu ngày và người thực hiện ở changelog — điền trước khi "
              "render lại.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
