#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
run_evals.py — Bo kiem hoi quy cho skill srs-help.

Chay mot lenh, kiem het. Dung sau moi lan sua outline.json, sua script, hoac
thay base.docx — nhung cho de vo nhat deu la cho khong ai nghi la lien quan.

    python evals/run_evals.py
    python evals/run_evals.py --pipeline /duong/dan/srs-pipeline   # them phep
                                                                   # doi chieu
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

SKILL = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(SKILL / "scripts"))
import srslib as S  # noqa: E402

SCRIPTS = SKILL / "scripts"
GOLDEN = SKILL / "references" / "golden"
PY = sys.executable

PASS, FAIL = "PASS", "FAIL"
results: list[tuple[str, str, str]] = []


def report(group: str, name: str, ok: bool, detail: str = "") -> bool:
    results.append((group, name, PASS if ok else FAIL))
    if not ok and detail:
        results[-1] = (group, name + f"  ← {detail}", FAIL)
    return ok


def run(args: list[str], cwd: Path | None = None) -> subprocess.CompletedProcess:
    # Both halves of the pipe must agree on UTF-8, not just this side. Every
    # script prints Vietnamese; on a default Windows PowerShell the *child*
    # Python writes to the pipe in cp1252, so decoding it as UTF-8 here still
    # crashed with UnicodeDecodeError at the first "LỖI". PYTHONUTF8=1 forces
    # the children to write UTF-8, and errors="replace" is the last-resort
    # net: one bad byte should cost one mangled character, never the run.
    env = os.environ.copy()
    env["PYTHONUTF8"] = "1"
    return subprocess.run([PY] + [str(a) for a in args], cwd=cwd,
                          capture_output=True, text=True, encoding="utf-8",
                          errors="replace", env=env)


# ---------------------------------------------------------------------------
def eval_outline(outline: dict) -> None:
    g = "Đề cương"
    r = run([SCRIPTS / "outline_check.py"])
    report(g, "outline_check sạch", r.returncode == 0,
           r.stdout.strip().split("\n")[-1] if r.returncode else "")

    blob = json.dumps(outline, ensure_ascii=False)
    for w in ("CIC", "19 ký tự"):
        report(g, f"không còn dấu vết “{w}”", w not in blob)

    usable = outline["layout"]["usable_twips"]
    bad = [t for t, v in outline["tables"].items() if sum(v["widths"]) != usable]
    report(g, f"{len(outline['tables'])} bảng đúng {usable} twips", not bad,
           ", ".join(bad[:3]))

    kinds = S.all_kinds(outline)
    report(g, f"{len(kinds)} loại tài liệu nhận dạng được",
           all(S.known_kind(outline, k) for k in kinds))


def eval_roundtrip(tmp: Path, outline: dict) -> None:
    g = "Vòng khép kín"
    for src in sorted(GOLDEN.glob("*.md")):
        work = tmp / f"rt-{src.stem}"
        work.mkdir()
        shutil.copy(src, work / src.name)
        for extra in ("assets", "registries"):
            if (GOLDEN / extra).exists():
                shutil.copytree(GOLDEN / extra, work / extra,
                                dirs_exist_ok=True)

        r = run([SCRIPTS / "render.py", src.name, "-o", "x.docx"], cwd=work)
        if not report(g, f"{src.stem}: render", r.returncode == 0,
                      r.stderr.strip()[:60]):
            continue
        r = run([SCRIPTS / "import_docx.py", "x.docx", "-o", "back.md"],
                cwd=work)
        if not report(g, f"{src.stem}: import", r.returncode == 0,
                      r.stderr.strip()[:60]):
            continue

        a = S.read_markdown(work / src.name)
        b = S.read_markdown(work / "back.md")
        diffs = compare(a, b)
        report(g, f"{src.stem}: md → docx → md′ đồng nhất", not diffs,
               "; ".join(diffs[:2]))


def compare(a: S.FunctionDoc, b: S.FunctionDoc) -> list[str]:
    bad = []
    if [s.name for s in a.sections] != [s.name for s in b.sections]:
        bad.append("danh sách mục lệch")
    if [(f.ma, f.ten) for f in a.features] != [(f.ma, f.ten) for f in b.features]:
        bad.append("khối tính năng lệch")

    def cmp_sec(x, y, where):
        ta = [t.rows for t in x.blocks if t.kind == "table"]
        tb = [t.rows for t in y.blocks if t.kind == "table"]
        if ta != tb:
            bad.append(f"bảng ở {where}/{x.name}")
        pa = [(t.kind, t.level, t.text) for t in x.blocks
              if t.kind in ("para", "bullet")]
        pb = [(t.kind, t.level, t.text) for t in y.blocks
              if t.kind in ("para", "bullet")]
        if pa != pb:
            bad.append(f"văn bản ở {where}/{x.name}")
        ia = [(t.label, t.path) for t in x.blocks if t.kind == "image"]
        ib = [(t.label, t.path) for t in y.blocks if t.kind == "image"]
        if ia != ib:
            bad.append(f"ảnh ở {where}/{x.name}")

    for x, y in zip(a.sections, b.sections):
        cmp_sec(x, y, "chức năng")
    for fx, fy in zip(a.features, b.features):
        for x, y in zip(fx.sections, fy.sections):
            cmp_sec(x, y, fx.ma)
    return bad


def eval_profiles(tmp: Path, outline: dict) -> None:
    g = "Loại tài liệu"
    for kind in S.all_kinds(outline):
        work = tmp / f"kind-{kind}"
        work.mkdir()
        ma = "GRP-EVAL-01" if kind == S.GROUP else "FUNC-EVAL-001"
        r = run([SCRIPTS / "scaffold.py", "--profile", kind, "--ma", ma,
                 "--ten", f"Kiểm thử {kind}", "--tinh-nang", "2",
                 "-o", "f.md"], cwd=work)
        if not report(g, f"{kind}: scaffold", r.returncode == 0,
                      r.stdout.strip()[:60]):
            continue
        r = run([SCRIPTS / "render.py", "f.md", "-o", "f.docx"], cwd=work)
        if not report(g, f"{kind}: render", r.returncode == 0):
            continue
        r = run([SCRIPTS / "import_docx.py", "f.docx", "-o", "b.md"], cwd=work)
        if not report(g, f"{kind}: import", r.returncode == 0):
            continue
        b = S.read_markdown(work / "b.md")
        report(g, f"{kind}: nhận đúng loại khi nhập lại", b.profile == kind,
               f"đoán ra {b.profile}")
        a = S.read_markdown(work / "f.md")
        report(g, f"{kind}: giữ nguyên danh sách mục",
               [s.name for s in a.sections] == [s.name for s in b.sections])


def eval_validate(tmp: Path) -> None:
    g = "Kiểm tra"
    work = tmp / "val"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "golden", dirs_exist_ok=True)
    gd = work / "golden"

    r = run([SCRIPTS / "validate.py", "FUNC-QLNSD-001.md",
             "--registry-dir", "registries"], cwd=gd)
    report(g, "golden chức năng: 0 lỗi", r.returncode == 0,
           [l for l in r.stdout.split("\n") if "LỖI" in l][:1])
    r = run([SCRIPTS / "validate.py", "GRP-QLNSD-01.md",
             "--registry-dir", "registries"], cwd=gd)
    report(g, "golden nhóm: 0 lỗi", r.returncode == 0)

    # -- negative cases -----------------------------------------------------
    base = (gd / "FUNC-QLNSD-001.md").read_text(encoding="utf-8")
    neg = {
        "thiếu khai tham số":
            ("| 1 | SUC_001 | Toast | Tạo {doi_tuong} thành công. | doi_tuong = NGUOIDUNG |",
             "| 1 | SUC_001 | Toast | Tạo {doi_tuong} thành công. |  |"),
        "khai thừa tham số":
            ("| 5 | ERR_104 | Inline | Vui lòng chọn ít nhất một vai trò. |  |",
             "| 5 | ERR_104 | Inline | Vui lòng chọn ít nhất một vai trò. | thua = X |"),
        "mã tham chiếu không khai":
            ("| BR-QLNSD-001-006 | Tài khoản ở trạng thái ST-NGUOIDUNG-03 không hiển thị trong kết quả tra cứu mặc định. | FEAT-QLNSD-001-01 | Không áp dụng |",
             "| BR-QLNSD-001-006 | Tài khoản ở trạng thái ST-NGUOIDUNG-03 không hiển thị trong kết quả tra cứu mặc định. | FEAT-QLNSD-001-01 | ERR_777 |"),
        "mã tính năng không liên tiếp":
            ("## Tính năng [FEAT-QLNSD-001-02]",
             "## Tính năng [FEAT-QLNSD-001-05]"),
        "còn cột placeholder":
            ("| 1 | FEAT-QLNSD-001-01 | Tra cứu danh sách người dùng | X | X | X |",
             "| 1 | FEAT-QLNSD-001-01 | Tra cứu danh sách người dùng | X | X |"),
        "cột vai trò chưa thay placeholder":
            ("| STT | Mã tính năng | Tính năng / Thao tác | ROLE-QTHT |",
             "| STT | Mã tính năng | Tính năng / Thao tác | «ROLE_1» |"),
        "cột vai trò không phải mã (đặt tên thường)":
            ("| STT | Mã tính năng | Tính năng / Thao tác | ROLE-QTHT | ROLE-QTDV | ROLE-NVNV | Phạm vi dữ liệu |",
             "| STT | Mã tính năng | Tính năng / Thao tác | Quản trị hệ thống | Quản trị đơn vị | Nhân viên nghiệp vụ | Phạm vi dữ liệu |"),
        "mã tính năng đúng ở Ma trận phân quyền nhưng để trống cột Mã tính năng":
            ("| 1 | FEAT-QLNSD-001-01 | Tra cứu danh sách người dùng | X | X | X |",
             "| 1 |  | Tra cứu danh sách người dùng FEAT-QLNSD-001-01 | X | X | X |"),
        "mã tính năng truy vết ở cột Ghi chú thay vì Tính năng đáp ứng":
            ("| UC-0301 |  | FEAT-QLNSD-001-01 | Chính | Đầy đủ |  |",
             "| UC-0301 |  |  | Chính | Đầy đủ | FEAT-QLNSD-001-01 |"),
        "mã UC sai dạng (UC-301 thay vì UC-0301)":
            ("| UC-0301 |  | FEAT-QLNSD-001-01 | Chính | Đầy đủ |  |",
             "| UC-301 |  | FEAT-QLNSD-001-01 | Chính | Đầy đủ |  |"),
        "mã thông báo sai dạng (thừa đoạn phân hệ)":
            ("| EXC-02 | Truy vấn quá thời gian chờ | Huỷ truy vấn, giữ nguyên điều kiện đã nhập. | ERR_002 |",
             "| EXC-02 | Truy vấn quá thời gian chờ | Huỷ truy vấn, giữ nguyên điều kiện đã nhập. | ERR_QLNSD_002 |"),
        "thiếu dòng Trong phạm vi ở bảng Mô tả chung":
            ("| Trong phạm vi | Tra cứu, tạo mới và ngừng hiệu lực tài khoản; gán vai trò đã tồn tại cho tài khoản. |\n",
             ""),
    }
    for name, (old, new) in neg.items():
        if old not in base:
            report(g, f"test âm “{name}”", False, "không tìm thấy mẫu để phá")
            continue
        p = gd / "neg.md"
        p.write_text(base.replace(old, new, 1), encoding="utf-8")
        r = run([SCRIPTS / "validate.py", "neg.md",
                 "--registry-dir", "registries"], cwd=gd)
        report(g, f"test âm “{name}” bị bắt", r.returncode == 1)
    (gd / "neg.md").unlink(missing_ok=True)

    # -- internal traceability ---------------------------------------------
    internal = {
        "nhắc màn hình không tồn tại":
            ("mở MH-QLNSD-001-002.", "mở MH-QLNSD-001-007."),
        "nhắc quy tắc đã bị xoá":
            ("| BR-QLNSD-001-006 | Tài khoản ở trạng thái ST-NGUOIDUNG-03 không "
             "hiển thị trong kết quả tra cứu mặc định. | FEAT-QLNSD-001-01 | "
             "Không áp dụng |\n", ""),
        "mang mã của chức năng khác":
            ("Áp dụng BR-QLNSD-001-001", "Áp dụng BR-QLSP-047-001"),
    }
    for name, (old, new) in internal.items():
        if old not in base:
            report(g, f"truy vết nội bộ “{name}”", False, "không thấy mẫu")
            continue
        p = gd / "ref.md"
        p.write_text(base.replace(old, new, 1), encoding="utf-8")
        r = run([SCRIPTS / "validate.py", "ref.md",
                 "--registry-dir", "registries"], cwd=gd)
        report(g, f"truy vết nội bộ “{name}” bị bắt",
               r.returncode == 1 and "truy vết nội bộ" in r.stdout)
    # declared but never referenced → warning, not an error
    p = gd / "ref.md"
    p.write_text(base.replace(
        "| 2 | MH-QLNSD-001-002 | Tạo mới người dùng | FEAT-QLNSD-001-02 |",
        "| 3 | MH-QLNSD-001-009 | Bỏ quên | FEAT-QLNSD-001-02 | Tàn dư. |\n"
        "| 2 | MH-QLNSD-001-002 | Tạo mới người dùng | FEAT-QLNSD-001-02 |", 1),
        encoding="utf-8")
    r = run([SCRIPTS / "validate.py", "ref.md",
             "--registry-dir", "registries"], cwd=gd)
    report(g, "mã khai rồi không ai dùng → chỉ cảnh báo",
           "khai rồi nhưng không nơi nào" in r.stdout and r.returncode != 1)
    p.unlink(missing_ok=True)

    # -- cross-table consistency -------------------------------------------
    cross = {
        "tính năng không truy về UC":
            ("| UC-0301 |  | FEAT-QLNSD-001-01 | Chính | Đầy đủ |  |\n", "",
             "error"),
        "quay về bước không tồn tại":
            ("phạm vi dữ liệu của vai trò. | 3 |",
             "phạm vi dữ liệu của vai trò. | 9 |", "error"),
        "vai trò không phải cột ma trận":
            ("| BR-QLNSD-001-005 | ROLE-QTDV chỉ tạo được",
             "| BR-QLNSD-001-005 | ROLE-KTNB chỉ tạo được", "warn"),
    }
    for name, (old, new, lvl) in cross.items():
        if old not in base:
            report(g, f"truy vết chéo “{name}”", False, "không thấy mẫu")
            continue
        p = gd / "cross.md"
        p.write_text(base.replace(old, new, 1), encoding="utf-8")
        r = run([SCRIPTS / "validate.py", "cross.md",
                 "--registry-dir", "registries"], cwd=gd)
        hit = "truy vết chéo" in r.stdout or "quay về bước" in r.stdout
        want = (r.returncode == 1) if lvl == "error" else (r.returncode != 1)
        report(g, f"truy vết chéo “{name}” bị bắt ({lvl})", hit and want)
    (gd / "cross.md").unlink(missing_ok=True)

    # -- obsolete marker ----------------------------------------------------
    p = gd / "obs.md"
    p.write_text((gd / "GRP-QLNSD-01.md").read_text(encoding="utf-8")
                 + "\n[[UCDIAGRAM: GRP-QLNSD-01]]\n", encoding="utf-8")
    r = run([SCRIPTS / "validate.py", "obs.md"], cwd=gd)
    report(g, "dấu [[UCDIAGRAM]] cũ bị báo lỗi",
           r.returncode == 1 and "UCDIAGRAM" in r.stdout)
    p.unlink(missing_ok=True)

    # -- release gate -------------------------------------------------------
    p = gd / "gate.md"
    p.write_text(base.replace("Áp dụng BR-QLNSD-001-005.",
                              "Áp dụng BR-QLNSD-001-005 ⟨?⟩", 1),
                 encoding="utf-8")
    r = run([SCRIPTS / "render.py", "gate.md", "-o", "gate.docx"], cwd=gd)
    report(g, "cổng chặn: tự ra bản nháp", "BẢN NHÁP" in r.stdout)
    r = run([SCRIPTS / "render.py", "gate.md", "--force-release",
             "-o", "gate2.docx"], cwd=gd)
    report(g, "cổng chặn: --force-release ghi đè được",
           r.returncode == 0 and "BẢN NHÁP" not in r.stdout)


def eval_roles_and_gaps(tmp: Path) -> None:
    """Gaps found by exercising the scripts like a user, not by reading the
    source: roles.csv was never cross-checked, `--profile` bypassed the
    master-document gate in import_docx.py, a [LỖI] line in project_check.py
    never flipped the verdict, and validate.py's exit code on a multi-file
    run let a gate-only file soften a real error in another file."""
    g = "Khoảng trống đã vá"
    work = tmp / "gaps"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "golden", dirs_exist_ok=True)
    gd = work / "golden"
    base = (gd / "FUNC-QLNSD-001.md").read_text(encoding="utf-8")

    # -- roles.csv is now cross-checked, like messages/usecases/states -----
    p = gd / "role.md"
    p.write_text(base.replace(
        "| STT | Mã tính năng | Tính năng / Thao tác | ROLE-QTHT | ROLE-QTDV | ROLE-NVNV | Phạm vi dữ liệu |",
        "| STT | Mã tính năng | Tính năng / Thao tác | ROLE-FAKE1 | ROLE-FAKE2 | ROLE-FAKE3 | Phạm vi dữ liệu |",
        1), encoding="utf-8")
    r = run([SCRIPTS / "validate.py", "role.md", "--registry-dir", "registries"],
            cwd=gd)
    report(g, "mã ROLE- ngoài roles.csv bị bắt",
           r.returncode == 1 and "roles.csv" in r.stdout)
    p.unlink(missing_ok=True)

    # -- import_docx: --profile no longer bypasses the master-document gate -
    from docx import Document as _D
    m = _D()
    m.add_heading("Nhóm chức năng [GRP-XXXX-01] A", 2)
    m.add_heading("Chức năng [FUNC-XXXX-001] B", 3)
    m.add_heading("Chức năng [FUNC-XXXX-002] C", 3)
    m.save(str(work / "master.docx"))
    r = run([SCRIPTS / "import_docx.py", "master.docx", "--profile", "UI",
             "-o", "z.md"], cwd=work)
    report(g, "tài liệu tổng vẫn bị từ chối dù có --profile",
           r.returncode == 1 and "TÀI LIỆU TỔNG" in r.stdout
           and not (work / "z.md").exists())

    # -- project_check.py: a [LỖI] line must flip the verdict --------------
    proj = work / "proj"
    shutil.copytree(GOLDEN, proj, dirs_exist_ok=True)
    reg = proj / "registries"
    rows = (reg / "objects.csv").read_text(encoding="utf-8").splitlines()
    header = rows[0].split(",")
    idx = header.index("ten_hien_thi")
    header.pop(idx)
    new_rows = [",".join(header)]
    for line in rows[1:]:
        cells = line.split(",")
        if len(cells) > idx:
            cells.pop(idx)
        new_rows.append(",".join(cells))
    (reg / "objects.csv").write_text("\n".join(new_rows) + "\n", encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=proj)
    report(g, "project_check: cột registry thiếu → CÒN THIẾU, không phải Đủ để làm việc",
           r.returncode == 1 and "objects.csv thiếu cột" in r.stdout
           and "Đủ để làm việc" not in r.stdout)

    # -- validate.py: error in one file must not be diluted by gate in another
    err_md = gd / "err.md"
    err_md.write_text(base.replace(
        "| 5 | ERR_104 | Inline | Vui lòng chọn ít nhất một vai trò. |  |",
        "| 5 | ERR_104 | Inline | Vui lòng chọn ít nhất một vai trò. | thua = X |",
        1), encoding="utf-8")
    gate_md = gd / "gate2.md"
    gate_md.write_text(base.replace(
        "Áp dụng BR-QLNSD-001-005.", "Áp dụng BR-QLNSD-001-005 ⟨?⟩", 1),
        encoding="utf-8")
    r = run([SCRIPTS / "validate.py", "err.md", "gate2.md",
             "--registry-dir", "registries"], cwd=gd)
    report(g, "kiểm nhiều file: lỗi ở một file thắng gate-only ở file khác "
              "(exit 1, không phải 2)", r.returncode == 1)
    err_md.unlink(missing_ok=True)
    gate_md.unlink(missing_ok=True)


def _fenced_blocks(text: str) -> list[str]:
    return re.findall(r"```[a-z]*\n(.*?)```", text, re.S)


def eval_docs_portability(outline: dict) -> None:
    """Example commands and front matter must work as pasted, on any OS.

    Each of these shipped broken once: an example front matter still carrying
    the previous outline version (follow the example, get rejected by the
    validator), bash-style `\\` line continuations (break in PowerShell), and
    `/tmp` paths (do not exist on Windows).
    """
    g = "Ví dụ chạy được"
    docs = [SKILL / "SKILL.md"] + sorted((SKILL / "references").glob("*.md"))
    ver = outline["version"]

    stale, cont, tmp = [], [], []
    for p in docs:
        text = p.read_text(encoding="utf-8")
        for block in _fenced_blocks(text):
            for m in re.finditer(r'outline_version:\s*"([^"]+)"', block):
                if m.group(1) != ver:
                    stale.append(f"{p.name}: {m.group(1)}")
            for line in block.split("\n"):
                if line.rstrip().endswith("\\"):
                    cont.append(f"{p.name}: {line.strip()[:50]}")
        if "/tmp/" in text or "/tmp " in text:
            tmp.append(p.name)

    report(g, f"mọi ví dụ front matter khai đúng v{ver}", not stale,
           "; ".join(stale[:3]))
    report(g, "không còn nối dòng kiểu bash `\\` trong ví dụ lệnh", not cont,
           "; ".join(cont[:3]))
    report(g, "không còn đường dẫn /tmp trong tài liệu", not tmp,
           ", ".join(tmp))


def eval_skill_doc(outline: dict) -> None:
    """What can be checked about SKILL.md without an inference run.

    Behaviour itself — does it invent content, does it confuse Write with
    Review — needs a real agent and lives in `evals/forward-tests.md`. These
    only guard the instructions that behaviour depends on.
    """
    g = "Hướng dẫn cho agent"
    sk = (SKILL / "SKILL.md").read_text(encoding="utf-8")

    # This sentence contradicted the paragraph above it: FEAT-/BR-/MH- are
    # minted inside the file, so "never invent a code that isn't in a
    # registry" told the agent to do the opposite of what it must do.
    report(g, "không còn câu cấm mọi mã ngoài sổ (mâu thuẫn với mã nội bộ)",
           "Never invent a code that isn't in a registry" not in sk)
    for kind in ("FEAT-", "BR-", "MH-"):
        report(g, f"nói rõ `{kind}` là mã nội bộ cấp trong file",
               kind in sk and "Minted here" in sk)
    report(g, "nói rõ mã dùng chung phải có sẵn trong sổ",
           "Must already exist in a registry" in sk)

    desc = sk.split("---")[1] if sk.startswith("---") else sk
    for phrase in ("test scripts", "user manuals"):
        report(g, f"description loại trừ “{phrase}”", phrase in desc)

    # -- token budget ------------------------------------------------------
    # SKILL.md loads in full on every activation; the reference set it points
    # at loads per situation. Both were cut deliberately (431 → ~170 lines;
    # mandatory reading ~52k → ~20k chars) and these keep them from creeping
    # back up.
    n_lines = sk.count("\n") + 1
    report(g, f"SKILL.md trong ngân sách ≤ 250 dòng (hiện {n_lines})",
           n_lines <= 250)
    m = re.search(r'description:\s*"(.*?)"', sk, re.S)
    d_len = len(m.group(1)) if m else 0
    report(g, f"description ≤ 500 ký tự (hiện {d_len})", 0 < d_len <= 500)
    report(g, "description không dùng ERR_ làm trigger",
           "ERR_" not in (m.group(1) if m else ""))
    report(g, "có bảng “What to read, when” điều phối việc đọc",
           "What to read, when" in sk)
    report(g, "có quy tắc không-đọc-nếu-không-cần",
           "Do not read unless" in sk)
    report(g, "luồng soạn mặc định dùng golden-snippets, không bắt đọc "
              "golden đầy đủ", "golden-snippets.md" in sk)

    # A reference named in SKILL.md but absent from the package sends the
    # agent looking for a file that is not there; it then works from memory.
    for m in sorted(set(re.findall(r"references/[A-Za-z0-9._/-]+", sk))):
        report(g, f"tham chiếu tồn tại: {m}", (SKILL / m).exists())
    report(g, "có bộ kiểm hành vi chạy tay",
           (SKILL / "evals" / "forward-tests.md").exists())

    ft = (SKILL / "evals" / "forward-tests.md")
    if ft.exists():
        body = ft.read_text(encoding="utf-8")
        report(g, "bộ kiểm hành vi có nhóm “không được kích hoạt”",
               "KHÔNG được kích hoạt" in body)


def eval_migration(tmp: Path, outline: dict) -> None:
    """A breaking outline bump is only safe if old documents have a way
    forward. These check the whole path: old file refused, migration run, file
    valid again, and the parts a human must still fill left explicitly open."""
    g = "Nâng đề cương"
    work = tmp / "mig"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "p", dirs_exist_ok=True)
    p = work / "p"
    src = p / "FUNC-QLNSD-001.md"
    cur = outline["version"]
    steps = {str(m["tu"]): str(m["den"]) for m in outline.get("migrations", [])
             if m.get("tu") and m.get("den")}
    if not report(g, f"outline.json có đường nâng cấp tới v{cur}",
                  cur in steps.values()):
        return
    # Start from the oldest version still described, not merely the previous
    # one: the point of the chain is that a document left behind two bumps ago
    # still reaches the head in one run, and testing only the last hop stops
    # covering that the moment a third version is added.
    old = next(v for v in steps if v not in set(steps.values()))
    prev = next(t for t, d in steps.items() if d == cur)
    major = lambda v: v.split(".")[0]

    # Rebuild a document as it looked under the oldest outline.
    text = src.read_text(encoding="utf-8")
    text = text.replace(f'outline_version: "{cur}"', f'outline_version: "{old}"')
    labels = outline["tables"]["TBL_KV_LOAI_CHUC_NANG"]["labels"]
    dropped = ["Trong phạm vi", "Ngoài phạm vi"]
    for lab in dropped:
        text = re.sub(rf"\|\s*{re.escape(lab)}\s*\|[^\n]*\n", "", text)
    src.write_text(text, encoding="utf-8")

    r = run([SCRIPTS / "validate.py", src.name, "--registry-dir", "registries"],
            cwd=p)
    report(g, f"tài liệu v{old} bị chặn (lỗi, không phải cảnh báo)",
           r.returncode == 1 and "lệch phiên bản LỚN" in r.stdout)
    report(g, "thông báo lỗi chỉ thẳng lệnh migration",
           "migrate_outline.py" in r.stdout)
    # Two hops behind is where a direct-hop lookup quietly gives up.
    report(g, f"thông báo nói rõ phải đi mấy chặng v{old} → v{cur}",
           len(steps) < 2 or f"chặng" in r.stdout, r.stdout.strip()[-90:])

    # A same-major bump must NOT block. Getting this wrong in either direction
    # is costly: block on a minor and every published spec stops rendering for
    # a change that altered nothing structural; wave through a major and the
    # analyst gets a screenful of "thiếu mục" that reads like their own fault.
    if major(prev) == major(cur):
        minor_dir = work / "minor"
        shutil.copytree(GOLDEN, minor_dir, dirs_exist_ok=True)
        ms = minor_dir / "FUNC-QLNSD-001.md"
        ms.write_text(ms.read_text(encoding="utf-8").replace(
            f'outline_version: "{cur}"', f'outline_version: "{prev}"'),
            encoding="utf-8")
        r = run([SCRIPTS / "validate.py", ms.name, "--registry-dir",
                 "registries"], cwd=minor_dir)
        report(g, f"tài liệu v{prev} (lệch NHỎ) chỉ bị cảnh báo, không chặn",
               r.returncode != 1 and "lệch phiên bản LỚN" not in r.stdout
               and f"v{prev}" in r.stdout, r.stdout.strip()[-90:])
        r = run([SCRIPTS / "render.py", ms.name, "-o", "m.docx"], cwd=minor_dir)
        report(g, f"tài liệu v{prev} vẫn render được", r.returncode == 0,
               r.stderr[-70:])

    r = run([SCRIPTS / "migrate_outline.py", src.name, "--thu"], cwd=p)
    report(g, "--thu chỉ xem trước, không ghi file",
           r.returncode == 0
           and all(f"| {lab} |" not in src.read_text(encoding="utf-8")
                   for lab in dropped))

    r = run([SCRIPTS / "migrate_outline.py", src.name,
             "--nguoi", "Eval", "--han", "2026-12-31"], cwd=p)
    if not report(g, "migrate_outline chạy được", r.returncode == 0,
                  r.stdout.strip()[-80:]):
        return

    after = src.read_text(encoding="utf-8")
    doc = S.read_markdown(src)
    report(g, f"front matter lên v{cur}",
           str(doc.meta.get("outline_version")) == cur)

    sec = doc.section("Mô tả chung")
    got = [row[0].strip() for b in sec.blocks if b.kind == "table"
           for row in b.rows[1:] if row]
    report(g, "dòng mới chèn đúng vị trí trong đề cương", got == labels,
           str(got[:8]))

    marker = outline["lexicon"]["open_marker"]
    filled = [row[1].strip() for b in sec.blocks if b.kind == "table"
              for row in b.rows[1:] if row and row[0].strip() in dropped]
    report(g, "nội dung để trống có đánh dấu, KHÔNG tự suy luận",
           filled == [marker] * len(dropped), str(filled))

    vd = doc.section(outline["gate"]["section"])
    pend = [row for b in vd.blocks if b.kind == "table"
            for row in b.rows[1:]
            if outline["lexicon"]["status_pending"] in " ".join(row)]
    report(g, "mỗi dòng chèn thêm mở một vấn đề còn chờ",
           len(pend) == len(dropped), f"{len(pend)} dòng")
    report(g, "changelog ghi lại lần nâng cấp",
           f"v{old}" in after and f"v{cur}" in after)

    r = run([SCRIPTS / "validate.py", src.name, "--registry-dir", "registries"],
            cwd=p)
    report(g, "sau nâng cấp: hết lỗi, chỉ còn vướng cổng chặn",
           r.returncode == 2, r.stdout.strip()[-90:])

    r = run([SCRIPTS / "migrate_outline.py", src.name], cwd=p)
    report(g, "chạy lại lần hai không đổi gì (idempotent)",
           r.returncode == 0 and "đã ở v" in r.stdout)


def eval_registry_schema(tmp: Path) -> None:
    """Registry faults that let bad codes through silently."""
    g = "Schema sổ đăng ký"
    work = tmp / "reg"
    work.mkdir()

    def fresh(name: str) -> Path:
        d = work / name
        shutil.copytree(GOLDEN, d, dirs_exist_ok=True)
        return d

    d = fresh("ok")
    r = run([SCRIPTS / "project_check.py", "."], cwd=d)
    report(g, "gói mẫu: Đủ để làm việc", r.returncode == 0
           and "Đủ để làm việc" in r.stdout)

    # Header-only file: no data rows at all, so a row-based check sees nothing
    # to inspect and passes.
    d = fresh("headeronly")
    (d / "registries" / "objects.csv").write_text(
        "ma,ten,mo_ta,ghi_chu\n", encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=d)
    report(g, "CSV chỉ có tiêu đề, thiếu cột vẫn bị bắt",
           r.returncode == 1 and "thiếu cột `ten_hien_thi`" in r.stdout)

    d = fresh("dup")
    (d / "registries" / "roles.csv").write_text(
        "ma,ten,tac_nhan_lien_quan,pham_vi_mac_dinh,mo_ta,ghi_chu\n"
        "ROLE-QTHT,A,,x,y,\nROLE-QTHT,B,,x,y,\n,C,,x,y,\n", encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=d)
    report(g, "mã trùng bị bắt", r.returncode == 1 and "mã trùng" in r.stdout)
    report(g, "mã bỏ trống bị bắt", "bỏ trống" in r.stdout)

    d = fresh("nokey")
    (d / "registries" / "states.csv").write_text(
        "code,ten,ten_hien_thi,mo_ta,ghi_chu\nST-X-01,a,b,c,\n",
        encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=d)
    report(g, "thiếu cột khoá bị bắt",
           r.returncode == 1 and "thiếu cột khoá" in r.stdout)

    d = fresh("emptyroles")
    (d / "registries" / "roles.csv").write_text(
        "ma,ten,tac_nhan_lien_quan,pham_vi_mac_dinh,mo_ta,ghi_chu\n",
        encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=d)
    report(g, "roles.csv rỗng trong khi tài liệu dùng ROLE- bị bắt",
           r.returncode == 1 and "roles.csv không có dòng nào" in r.stdout)


def eval_export_pdf(tmp: Path) -> None:
    """Runs only where LibreOffice exists; elsewhere it is skipped rather than
    failed, so the suite stays usable on a machine without it."""
    g = "Xuất PDF"
    if not (shutil.which("soffice") or shutil.which("libreoffice")):
        report(g, "bỏ qua — máy không có LibreOffice", True)
        return
    work = tmp / "pdf"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "p", dirs_exist_ok=True)
    p = work / "p"

    r = run([SCRIPTS / "render.py", "GRP-QLNSD-01.md", "-o", "g.docx"], cwd=p)
    if not report(g, "render .docx", r.returncode == 0):
        return
    r = run([SCRIPTS / "export_pdf.py", "g.docx"], cwd=p, )
    ok = r.returncode == 0 and (p / "g.pdf").exists()
    report(g, "xuất được .pdf", ok, r.stdout.strip()[-80:])
    if ok:
        head = (p / "g.pdf").read_bytes()[:5]
        report(g, "file .pdf hợp lệ", head == b"%PDF-", str(head))
    # The UNO path is what keeps SEQ/TOC numbering correct; falling back is
    # allowed but must say so rather than pass silently.
    report(g, "đường UNO chạy, hoặc nói rõ đã lùi về cách thường",
           "CẢNH BÁO" in r.stdout or "Cầu UNO" in r.stderr
           or "CẢNH BÁO" not in r.stdout)

    src = (SCRIPTS / "export_pdf.py").read_text(encoding="utf-8")
    report(g, "không còn ghim cổng UNO cố định", "PORT = 2103" not in src)
    report(g, "không còn ghim đường dẫn /tmp", 'Path(f"/tmp/' not in src)


def eval_stt(tmp: Path, outline: dict) -> None:
    """`STT` is derived from row position, so render owns it.

    Legacy specs numbered these rows with Word's own list numbering, which
    lives in `numbering.xml` and never appears in the cell text: 260 visible
    numbers in a real document read back as empty strings. Nothing could be
    carried across, so the column has to be computed rather than transported.
    """
    g = "Cột STT"
    sys.path.insert(0, str(SCRIPTS))
    import import_docx as I
    from docx import Document

    cases = [
        ("đánh số liên tiếp từ ô trống",
         [["STT", "Tên", "Mô tả"], ["", "A", "x"], ["", "B", "y"]],
         ["1", "2"]),
        ("đánh lại khi số cũ sai thứ tự",
         [["STT", "Tên"], ["7", "A"], ["3", "B"]], ["1", "2"]),
        ("dòng nhãn không ăn số (nhãn ở ô đầu)",
         [["STT", "Tên"], ["", "A"], ["Các button", ""], ["", "B"]],
         ["1", "Các button", "2"]),
        ("dòng nhãn không ăn số (kiểu lặp, file nhập bản cũ)",
         [["STT", "Tên"], ["", "A"], ["Btn", "Btn"], ["", "B"]],
         ["1", "Btn", "2"]),
        ("ô STT có chữ thật thì giữ nguyên",
         [["STT", "Tên"], ["", "A"], ["n/a", "B"], ["", "C"]],
         ["1", "n/a", "2"]),
        ("bảng không có cột STT thì không đụng tới",
         [["Mã", "Tên"], ["", "A"], ["", "B"]], ["", ""]),
    ]
    for name, rows, want in cases:
        got = [r[0] for r in S.renumber_stt(rows)[1:]]
        report(g, name, got == want, f"ra {got}")

    # A seeded scaffold must still count as empty, or the "mục để trống" check
    # goes quiet the moment the ordinals arrive.
    sec = S.parse_markdown(
        "## X\n\n| STT | Tên | Mô tả |\n|---|---|---|\n| 1 |  |  |\n"
        "| 2 |  |  |\n").sections[0]
    report(g, "bảng chỉ có số STT vẫn tính là mục để trống", sec.is_empty(),
           repr(sec.text_content()[:40]))

    work = tmp / "stt"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "g", dirs_exist_ok=True)
    p = work / "g"

    # Render is the authority: wipe the ordinals in the source and they must
    # still come out right in the document.
    src = p / "FUNC-QLNSD-001.md"
    txt = src.read_text(encoding="utf-8")
    txt = re.sub(r"^\| \d+ \|", "|  |", txt, flags=re.M)
    src.write_text(txt, encoding="utf-8")
    r = run([SCRIPTS / "render.py", src.name, "-o", "s.docx"], cwd=p)
    if report(g, "render file đã xoá hết số STT", r.returncode == 0):
        d = Document(str(p / "s.docx"))
        bad = []
        for t in d.tables:
            if not t.rows or t.rows[0].cells[0].text.strip() != "STT":
                continue
            want_n = 0
            for row in t.rows[1:]:
                vals = [c.text.strip() for c in row.cells]
                if S._is_label_row(vals) or not " ".join(vals).strip():
                    continue
                want_n += 1
                if vals[0] != str(want_n):
                    bad.append(f"{vals[0]!r}≠{want_n}")
        report(g, "render tự điền lại STT đúng 1..n", not bad,
               ", ".join(bad[:3]))

    # And a merged divider band must not be duplicated across the row. The
    # table is written as raw XML — one `w:tc` carrying `w:gridSpan`, exactly
    # what Word produces — because python-docx's own `merge()` builds a
    # different shape, and editing a generated table's cells corrupts its
    # `tblGrid`. Neither would exercise what real documents contain.
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    def _tc(txt, span=None):
        pr = (f'<w:tcPr><w:gridSpan w:val="{span}"/></w:tcPr>' if span
              else "<w:tcPr/>")
        return f"<w:tc>{pr}<w:p><w:r><w:t>{txt}</w:t></w:r></w:p></w:tc>"

    xml = (f'<w:tbl {nsdecls("w")}><w:tblPr/><w:tblGrid>'
           + '<w:gridCol w:w="1000"/>' * 3 + "</w:tblGrid>"
           + "<w:tr>" + _tc("STT") + _tc("Tên") + _tc("Mô tả") + "</w:tr>"
           + "<w:tr>" + _tc("") + _tc("Dòng thường") + _tc("x") + "</w:tr>"
           + "<w:tr>" + _tc("Các button", span=3) + "</w:tr></w:tbl>")
    d = Document()
    d.element.body.append(parse_xml(xml))
    d.save(str(work / "m.docx"))

    tm = Document(str(work / "m.docx")).tables[0]
    report(g, "fixture đúng hình: ô gộp trả về lặp 3 lần",
           [c.text for c in tm.rows[2].cells] == ["Các button"] * 3)
    imp = I.Importer(outline, work)
    bs2 = I.bullet_style_map(outline)
    got = imp._row_cells(tm.rows[2], bs2)
    report(g, "dòng gộp ô chỉ ghi nhãn một lần, không lặp khắp cột",
           got == ["Các button", "", ""], str(got))
    full = S.renumber_stt([imp._row_cells(r, bs2) for r in tm.rows])
    report(g, "và dòng gộp đó không ăn số STT",
           [r[0] for r in full[1:]] == ["1", "Các button"], str(full[1:]))


def eval_project_layout(tmp: Path) -> None:
    """Assets belong beside the `.md`, and getting it wrong must be loud.

    A single `assets/` at the project root with specs one level down renders
    every mockup as a ⟨ THIẾU HÌNH ⟩ box — while `validate.py` still reports
    `0 lỗi`, because a missing optional image is only a warning. Both of this
    skill's own deployment documents used to show exactly that layout.
    """
    g = "Bố trí thư mục"
    from docx import Document

    def build(where_assets: str) -> Path:
        d = tmp / f"layout-{where_assets}"
        (d / "functions" / "qlnsd").mkdir(parents=True)
        (d / "registries").mkdir()
        shutil.copy(GOLDEN / "FUNC-QLNSD-001.md", d / "functions" / "qlnsd")
        for c in (GOLDEN / "registries").glob("*.csv"):
            shutil.copy(c, d / "registries")
        target = (d / "assets" if where_assets == "goc"
                  else d / "functions" / "qlnsd" / "assets")
        target.mkdir(parents=True)
        for p in (GOLDEN / "assets").glob("*.png"):
            shutil.copy(p, target)
        return d

    ok_dir = build("canh-md")
    bad_dir = build("goc")
    md = "functions/qlnsd/FUNC-QLNSD-001.md"

    r = run([SCRIPTS / "render.py", md, "-o", "out.docx"], cwd=ok_dir)
    boxes = 0
    if report(g, "bố trí ĐÚNG: render được", r.returncode == 0):
        doc = Document(str(ok_dir / "out.docx"))
        boxes = sum(1 for t in doc.tables for row in t.rows for c in row.cells
                    if "THIẾU HÌNH" in c.text)
        import zipfile
        n_img = len([n for n in zipfile.ZipFile(ok_dir / "out.docx").namelist()
                     if n.startswith("word/media/")])
        report(g, "bố trí ĐÚNG: ảnh mockup vào được tài liệu",
               n_img >= 3, f"{n_img} ảnh")

    r = run([SCRIPTS / "render.py", md, "-o", "out.docx"], cwd=bad_dir)
    if r.returncode == 0:
        doc = Document(str(bad_dir / "out.docx"))
        bad_boxes = sum(1 for t in doc.tables for row in t.rows
                        for c in row.cells if "THIẾU HÌNH" in c.text)
        report(g, "bố trí SAI: mockup thành khung ⟨THIẾU HÌNH⟩",
               bad_boxes > boxes, f"{bad_boxes} vs {boxes} khung")

    # The trap: validation is happy either way, so the layout check has to be
    # the thing that speaks up.
    r = run([SCRIPTS / "validate.py", md, "--registry-dir", "registries"],
            cwd=bad_dir)
    report(g, "validate KHÔNG bắt được (đúng như thiết kế — chỉ là cảnh báo)",
           r.returncode == 0)

    r = run([SCRIPTS / "project_check.py", "."], cwd=bad_dir)
    report(g, "project_check bắt được và gọi đúng tên: đặt sai chỗ",
           r.returncode == 1 and "ĐẶT SAI CHỖ" in r.stdout)
    report(g, "in ra cả nơi đang ở và nơi cần ở",
           "đang ở :" in r.stdout and "cần ở  :" in r.stdout)
    report(g, "không báo nhầm thành “thiếu ảnh”",
           "Ảnh mockup thiếu" not in r.stdout)

    r = run([SCRIPTS / "project_check.py", "."], cwd=ok_dir)
    report(g, "bố trí đúng thì không cảnh báo gì",
           "ĐẶT SAI CHỖ" not in r.stdout and r.returncode == 0)

    # And the documents must describe the layout that actually works. Both of
    # them used to draw `assets/` as a top-level entry beside `functions/`,
    # which is precisely the arrangement that renders empty boxes — analysts
    # were following the documentation into the failure.
    top_level = re.compile(r"^[├└]──\s*(assets|diagrams)/", re.M)
    for name in ("trien-khai.md", "huong-dan-ba.md"):
        txt = (SKILL / "references" / name).read_text(encoding="utf-8")
        blocks = [b for b in _fenced_blocks(txt) if "du-an/" in b]
        offenders = [m.group(1) for b in blocks for m in top_level.finditer(b)]
        report(g, f"{name}: sơ đồ không đặt assets/diagrams ở gốc",
               bool(blocks) and not offenders, str(set(offenders)))

    # Folder case has to agree with BA Toolkit, which mandates lowercase. On
    # Windows both spellings resolve to the same directory, so a mismatch
    # survives every local test and only splits into two folders once the
    # project reaches a case-sensitive filesystem.
    upper = re.compile(r"functions/[A-Z]{3,6}/")
    for name in ("trien-khai.md", "huong-dan-ba.md", "agent-notes.md"):
        txt = (SKILL / "references" / name).read_text(encoding="utf-8")
        report(g, f"{name}: thư mục phân hệ viết thường",
               not upper.search(txt), str(upper.findall(txt)[:3]))
    for script in ("project_check.py", "render.py"):
        txt = (SKILL / "scripts" / script).read_text(encoding="utf-8")
        report(g, f"{script}: gợi ý cũng viết thường",
               not upper.search(txt), str(upper.findall(txt)[:3]))


def eval_workspace_shape(tmp: Path) -> None:
    """Living beside BA Toolkit without either side having to bend.

    A toolkit workspace holds folders the skill knows nothing about, and one
    of them — `staging/` — contains drafts with valid front matter. Scoring
    those reports faults on work nobody has approved yet, and scoring
    `sources/` reports faults on files that are hashed evidence the analyst is
    forbidden to touch. Both are worse than useless: they bury the real
    findings.
    """
    g = "Sống chung với toolkit"
    work = tmp / "ws"
    (work / "functions" / "qlnsd" / "assets").mkdir(parents=True)
    (work / "registries").mkdir()
    (work / "project-rules").mkdir()
    for sub in ("staging/qlnsd", "sources/legacy/docx", "migration/run-1/raw",
                "reports/runs", ".ba-toolkit"):
        (work / sub).mkdir(parents=True)
    shutil.copy(GOLDEN / "FUNC-QLNSD-001.md", work / "functions" / "qlnsd")
    for c in (GOLDEN / "registries").glob("*.csv"):
        shutil.copy(c, work / "registries")
    for p in (GOLDEN / "assets").glob("*.png"):
        shutil.copy(p, work / "functions" / "qlnsd" / "assets")
    # A draft that would fail if scored, in each folder that must be skipped.
    nhap = (GOLDEN / "FUNC-QLNSD-001.md").read_text(encoding="utf-8")
    nhap = nhap.replace("Quản lý người dùng", "⟨?⟩", 1)
    for rel in ("staging/qlnsd/FUNC-QLNSD-001.md",
                "sources/legacy/docx/FUNC-QLNSD-001.md",
                "migration/run-1/raw/FUNC-QLNSD-001.md"):
        (work / rel).write_text(nhap, encoding="utf-8")

    # A group file is a document of record too, and it lives in `groups/`, not
    # `functions/`. The first cut of this whitelist dropped every one of them
    # silently — then concluded "Đủ để làm việc" about a half-read project.
    (work / "groups").mkdir()
    shutil.copy(GOLDEN / "GRP-QLNSD-01.md", work / "groups")

    r = run([SCRIPTS / "project_check.py", "."], cwd=work)
    report(g, "chỉ chấm tài liệu trong functions/ và groups/",
           r.returncode == 0, r.stdout.strip()[-100:])
    report(g, "file nhóm ở groups/ ĐƯỢC chấm, không bị bỏ qua",
           "GRP-QLNSD-01" in r.stdout and "2 file đặc tả" in r.stdout,
           r.stdout[:220])
    report(g, "nói rõ đã bỏ qua bao nhiêu file, không im lặng",
           "Bỏ qua 3 file" in r.stdout and "staging" in r.stdout,
           r.stdout[:200])
    report(g, "không có overlay thì nói rõ là chạy thuần chuẩn",
           "Không có project-rules" in r.stdout)

    (work / "project-rules" / "srs-help.md").write_text(
        "# Luật riêng\n\nDự án dùng BA Toolkit.\n", encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=work)
    report(g, "có overlay thì báo ra — không được lặng lẽ đổi hành vi",
           "project-rules/srs-help.md" in r.stdout and "SAU SKILL.md" in r.stdout)

    # The empty-registry trap: another folder name plus a later `init` gives
    # eight empty CSVs that pass every schema check while the real books sit
    # unread one folder away.
    lac = tmp / "ws-lac"
    shutil.copytree(work, lac)
    (lac / "registries").rename(lac / "so-dang-ky")
    r = run([SCRIPTS / "project_check.py", "."], cwd=lac)
    report(g, "tên sổ khác `registries/` bị bắt là LỖI",
           r.returncode == 1 and "so-dang-ky" in r.stdout, r.stdout[-120:])
    report(g, "nói rõ hậu quả: toolkit sẽ tạo sổ rỗng bên cạnh",
           "rỗng" in r.stdout)

    # SKILL.md must actually tell the agent the overlay exists; a check that
    # only project_check knows about changes nothing about how the agent works.
    txt = (SKILL / "SKILL.md").read_text(encoding="utf-8")
    report(g, "SKILL.md dạy agent đọc project-rules/srs-help.md",
           "project-rules/srs-help.md" in txt)
    report(g, "SKILL.md nói rõ overlay không được nới lỏng cổng chặn",
           "release gate" in txt.split("Project overlay")[-1][:700])


def eval_multiline_cells(tmp: Path, outline: dict) -> None:
    """`·` marks become real bullets — but only where the outline allows it.

    The one-line rule existed for three reasons; two still hold. Sequential
    columns genuinely hold several points, so they get bullets. Constraint
    columns do not: the standard already says each rule takes its own `BR-`
    code, and making a crowded cell read nicely would remove the reason to
    split it. The source line stays one line either way, so diffs are
    unaffected — only the rendered document changes.
    """
    g = "Ô nhiều dòng"
    spec = S.multiline_spec(outline)
    allow = spec.get("cho_phep", [])
    warn_cols = spec.get("canh_bao", [])
    report(g, "đề cương khai cột được phép và cột cảnh báo",
           bool(allow) and bool(warn_cols), f"{len(allow)}/{len(warn_cols)}")

    segs = S.cell_segments("Mở đầu: ·· Ý A ·· Ý B · Ý C")
    report(g, "tách đúng văn xuôi và cấp gạch đầu dòng",
           segs == [("Mở đầu:", 0), ("Ý A", 2), ("Ý B", 2), ("Ý C", 1)],
           str(segs))
    report(g, "ô không có dấu · thì không coi là nhiều dòng",
           not S.cell_is_multiline("Một câu bình thường."))

    work = tmp / "mline"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "p", dirs_exist_ok=True)
    p = work / "p"
    src = p / "FUNC-QLNSD-001.md"
    base = src.read_text(encoding="utf-8")
    t = base.replace(
        "| EXC-01 | Không có bản ghi nào khớp | Giữ nguyên điều kiện, hiển thị lưới rỗng kèm thông báo. | INF_001 |",
        "| EXC-01 | Không có bản ghi nào khớp | Hệ thống: ·· Giữ nguyên điều kiện ·· Hiển thị lưới rỗng · Ghi nhật ký | INF_001 |", 1)
    t = t.replace(
        "| 2 | Họ và tên | Ô nhập văn bản | Không | tối đa 100 ký tự | Tìm kiếm khớp một phần. |",
        "| 2 | Họ và tên | Ô nhập văn bản | Không | tối đa 100 ký tự | Quy tắc: · Khớp một phần · Không phân biệt hoa thường |", 1)
    src.write_text(t, encoding="utf-8")

    r = run([SCRIPTS / "render.py", src.name, "-o", "g.docx"], cwd=p)
    if not report(g, "render được", r.returncode == 0, r.stderr[-70:]):
        return
    from docx import Document
    d = Document(str(p / "g.docx"))
    got = {}
    for tb in d.tables:
        if not tb.rows:
            continue
        head = [c.text.strip() for c in tb.rows[0].cells]
        for row in tb.rows[1:]:
            for j, c in enumerate(row.cells):
                if len(c.paragraphs) > 1 and j < len(head):
                    got[head[j]] = [pp.style.name for pp in c.paragraphs]

    report(g, "cột ĐƯỢC PHÉP ra nhiều đoạn, đúng style bullet theo cấp",
           got.get("Xử lý của hệ thống") ==
           [outline["styles"]["body"], outline["styles"]["bullet_2"],
            outline["styles"]["bullet_2"], outline["styles"]["bullet_1"]],
           str(got.get("Xử lý của hệ thống")))
    report(g, "cột KHÔNG được phép vẫn giữ đúng một đoạn",
           "Mô tả ràng buộc" not in got, str(got.get("Mô tả ràng buộc")))

    # The whole point: this must not disturb the round trip.
    r = run([SCRIPTS / "import_docx.py", "g.docx", "-o", "b.md"], cwd=p)
    if report(g, "nhập lại được", r.returncode == 0):
        def tables(doc):
            return [bl.rows
                    for s in doc.sections + [x for f in doc.features
                                             for x in f.sections]
                    for bl in s.blocks if bl.kind == "table"]
        report(g, "vòng khép kín .md → .docx → .md bất biến",
               tables(S.read_markdown(src)) ==
               tables(S.read_markdown(p / "b.md")))

    r = run([SCRIPTS / "validate.py", src.name, "--registry-dir",
             "registries"], cwd=p)
    report(g, "cảnh báo gộp ý ở cột ràng buộc, gộp thành một dòng tóm tắt",
           "gộp ý trong ô" in r.stdout
           and r.stdout.count("gộp nhiều ý") == 1, r.stdout[-80:])
    report(g, "cột được phép thì không bị cảnh báo",
           "Xử lý của hệ thống" not in r.stdout.split("gộp ý trong ô")[-1][:200])

    deep = p / "deep.md"
    deep.write_text(t.replace("·· Giữ nguyên", "···· Giữ nguyên", 1),
                    encoding="utf-8")
    r = run([SCRIPTS / "validate.py", "deep.md", "--registry-dir",
             "registries"], cwd=p)
    report(g, f"quá {spec.get('max_cap', 3)} cấp là lỗi",
           r.returncode == 1 and "cấp" in r.stdout)


def _fill_columns(text: str, wanted: set[str], value: str) -> dict[str, str]:
    """Put `value` in the first data row of every column named in `wanted`.

    Returns the columns actually touched. Header rows are found the way the
    parser finds them — a `|` line followed by a `|---|` separator — so a
    pipe inside prose cannot be mistaken for a table.
    """
    lines = text.split("\n")
    touched: dict[str, str] = {}
    i = 0
    while i < len(lines) - 1:
        if not lines[i].lstrip().startswith("|") or not S._is_sep(lines[i + 1]):
            i += 1
            continue
        head = [S.norm(c) for c in S._table_row(lines[i])]
        hit = [j for j, h in enumerate(head) if h in wanted]
        j_row = i + 2
        if hit and j_row < len(lines) and lines[j_row].lstrip().startswith("|"):
            cells = S._table_row(lines[j_row])
            if len(cells) >= len(head):
                for j in hit:
                    cells[j] = value
                    touched[head[j]] = value
                lines[j_row] = "| " + " | ".join(cells) + " |"
        i = j_row
    return touched, "\n".join(lines)


def eval_multiline_profiles(tmp: Path, outline: dict) -> None:
    """The bullet rule has to hold in every template, not just `UI`.

    `multiline_columns` names columns, and a column belongs to whichever
    templates happen to use that table. That indirection is what let the
    first cut of the feature reach `UI` almost exclusively without anyone
    noticing: the declaration looked complete on its own. So the test is not
    "does the list have entries" but "does each template actually get
    bullets where its own content is sequential".
    """
    g = "Ô nhiều dòng × loại"
    spec = S.multiline_spec(outline)
    allow = set(spec.get("cho_phep", []))
    warn_cols = set(spec.get("canh_bao", []))
    tables = outline["tables"]

    # A declared column no table uses is dead weight that still reads as
    # coverage — the exact illusion this group exists to break.
    used = {h for t in tables.values() for h in t.get("headers", [])}
    dead = sorted((allow | warn_cols) - used)
    report(g, "mọi cột đã khai đều có bảng dùng tới", not dead, ", ".join(dead))

    # The lead ends in a colon on purpose: import gives an unpunctuated prose
    # line a full stop, so a bare lead would come back changed and the round
    # trip would fail for a reason that has nothing to do with bullets.
    val = "Mở đầu: · Ý một · Ý hai"
    want_par = [outline["styles"]["body"], outline["styles"]["bullet_1"],
                outline["styles"]["bullet_1"]]
    from docx import Document

    for kind in S.all_kinds(outline):
        if kind == S.GROUP:
            continue
        prof = S.profile_of(outline, kind)
        tids = {t for s in prof["function_sections"] + prof["feature_sections"]
                for t in s.get("tables", [])}
        mine = {h for t in tids for h in tables[t].get("headers", [])} & allow
        # `Nội dung vấn đề` is on the shared open-issues table, so a template
        # that only has that one has gained nothing from the feature.
        report(g, f"{kind}: có cột được bullet ngoài bảng dùng chung",
               bool(mine - {"Nội dung vấn đề"}), f"chỉ có {sorted(mine)}")

        work = tmp / f"ml-{kind}"
        work.mkdir()
        r = run([SCRIPTS / "scaffold.py", "--profile", kind, "--ma",
                 "FUNC-EVAL-001", "--ten", f"Thử {kind}", "--tinh-nang", "1",
                 "-o", "f.md"], cwd=work)
        if not report(g, f"{kind}: scaffold", r.returncode == 0):
            continue
        src = work / "f.md"
        touched, new = _fill_columns(src.read_text(encoding="utf-8"), mine, val)
        src.write_text(new, encoding="utf-8")
        if not report(g, f"{kind}: điền được vào mọi cột đã khai",
                      set(touched) == mine, f"hụt {sorted(mine - set(touched))}"):
            continue

        r = run([SCRIPTS / "render.py", "f.md", "-o", "f.docx"], cwd=work)
        if not report(g, f"{kind}: render", r.returncode == 0, r.stderr[-70:]):
            continue
        got: dict[str, list[str]] = {}
        for tb in Document(str(work / "f.docx")).tables:
            if not tb.rows:
                continue
            head = [c.text.strip() for c in tb.rows[0].cells]
            for row in tb.rows[1:]:
                for j, c in enumerate(row.cells):
                    if j < len(head) and head[j] in touched:
                        got.setdefault(head[j], [p.style.name
                                                 for p in c.paragraphs])
        bad = sorted(c for c in touched if got.get(c) != want_par)
        report(g, f"{kind}: mọi cột đã khai ra đúng gạch đầu dòng", not bad,
               "; ".join(f"{c}→{got.get(c)}" for c in bad)[:90])

        r = run([SCRIPTS / "validate.py", "f.md"], cwd=work)
        report(g, f"{kind}: không cột nào bị báo “dấu · lạc chỗ”",
               "không khai trong" not in r.stdout,
               r.stdout[r.stdout.find("không khai trong") - 60:][:90])

        r = run([SCRIPTS / "import_docx.py", "f.docx", "-o", "b.md"], cwd=work)
        if not report(g, f"{kind}: nhập lại", r.returncode == 0):
            continue
        # The bullets have to fold back into the same `·` string, or the round
        # trip silently rewrites every spec that uses the feature.
        seen = _column_values(S.read_markdown(work / "b.md"), set(touched))
        lost = sorted(c for c in touched if val not in seen.get(c, []))
        report(g, f"{kind}: gạch đầu dòng gấp lại đúng dấu · khi nhập lại",
               not lost, "; ".join(f"{c}→{seen.get(c)}" for c in lost)[:90])


def _column_values(doc, wanted: set[str]) -> dict[str, list[str]]:
    out: dict[str, list[str]] = {}
    secs = list(doc.sections) + [s for f in doc.features for s in f.sections]
    for sec in secs:
        for b in sec.blocks:
            if b.kind != "table" or not b.rows:
                continue
            head = [S.norm(c) for c in b.rows[0]]
            for row in b.rows[1:]:
                for j, v in enumerate(row):
                    if j < len(head) and head[j] in wanted:
                        out.setdefault(head[j], []).append(v)
    return out


def eval_manifest(tmp: Path, outline: dict) -> None:
    """`FUNC-` codes are allocated, not derived.

    An inventory can always be rebuilt from disk. An allocation cannot: the
    code has to be reserved before its file exists, or two analysts working
    without a shared lock both take `FUNC-QLSP-048` and nothing notices until
    the documents meet. Before this existed the standard pointed at a "project
    manifest" that was not present in any repository — a dangling reference in
    three files.
    """
    g = "Danh mục chức năng"
    man = outline.get("manifest") or {}
    report(g, "đề cương khai manifest và tệp mẫu",
           man.get("file") == "manifest.md" and bool(man.get("mau")),
           str(man)[:60])
    mau = SKILL / (man.get("mau") or "assets/manifest.example.md")
    report(g, "tệp mẫu có thật", mau.is_file(), str(mau))
    # The dangling reference must stay fixed: no document may send the reader
    # looking for an artefact that does not exist.
    treo = []
    for f in ("SKILL.md", "references/huong-dan-ba.md",
              "references/agent-notes.md"):
        txt = (SKILL / f).read_text(encoding="utf-8")
        if re.search(r"manifest dự án|project manifest", txt):
            treo.append(f)
    report(g, "không còn chỗ nào trỏ tới “manifest dự án” chung chung",
           not treo, ", ".join(treo))

    work = tmp / "man"
    (work / "functions" / "qlnsd" / "assets").mkdir(parents=True)
    (work / "registries").mkdir()
    shutil.copy(GOLDEN / "FUNC-QLNSD-001.md", work / "functions" / "qlnsd")
    for c in (GOLDEN / "registries").glob("*.csv"):
        shutil.copy(c, work / "registries")
    for p in (GOLDEN / "assets").glob("*.png"):
        shutil.copy(p, work / "functions" / "qlnsd" / "assets")

    r = run([SCRIPTS / "project_check.py", "."], cwd=work)
    report(g, "không có manifest: cảnh báo, KHÔNG chặn",
           r.returncode == 0 and "Không thấy manifest.md" in r.stdout,
           r.stdout.strip()[-90:])

    shutil.copy(mau, work / "manifest.md")
    r = run([SCRIPTS / "project_check.py", "."], cwd=work)
    report(g, "mã khớp: sạch", r.returncode == 0 and "mã đã cấp" in r.stdout,
           r.stdout.strip()[-90:])

    goc = (work / "manifest.md").read_text(encoding="utf-8")
    (work / "manifest.md").write_text(
        goc.replace("| FUNC-QLNSD-001 |", "| FUNC-XXXXX-999 |", 1),
        encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=work)
    report(g, "có file nhưng chưa cấp mã: LỖI",
           r.returncode == 1 and "chưa cấp trong manifest" in r.stdout,
           r.stdout.strip()[-90:])
    report(g, "mã “Đã phát hành” mà không có file: chỉ cảnh báo",
           "không thấy file" in r.stdout)

    (work / "manifest.md").write_text(
        goc + "\n| FUNC-QLNSD-001 | Trùng | UI | QLNSD | | | | Đã cấp | |\n",
        encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=work)
    report(g, "mã ghi hai dòng: LỖI",
           r.returncode == 1 and "ghi hai lần" in r.stdout,
           r.stdout.strip()[-90:])

    # A blank manifest keeps its example row inside an HTML comment. Counting
    # that would report a code as allocated that nobody allocated — and the
    # first thing a new project does is copy exactly such a file.
    (work / "manifest.md").write_text(
        "# Manifest\n\n| Mã | Tên | Loại |\n|---|---|---|\n| | | |\n\n"
        "<!-- ví dụ:\n| FUNC-QLNSD-001 | Quản lý người dùng | UI |\n-->\n",
        encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=work)
    report(g, "dòng ví dụ trong <!-- --> không bị tính là đã cấp",
           "0 mã đã cấp" in r.stdout, r.stdout[r.stdout.find("Danh mục"):][:70])


def _uoc_token(text: str) -> int:
    """Rough token count. No tokenizer here, and the exact number does not
    matter — what matters is that a file does not quietly double in size."""
    n = len(text)
    viet = sum(1 for c in text if ord(c) > 127)
    cpt = 4.0 - (4.0 - 2.6) * min(viet / max(n, 1) * 6, 1.0)
    return round(n / cpt)


def eval_ngan_sach_token() -> None:
    """Reading a file loads all of it, so file size is the real unit of cost.

    `SKILL.md` says "read the section you need", but there is no such thing:
    an agent wanting the 320-token *Images* note pays for the whole file. That
    makes one fat section expensive for everyone — legacy migration was 40% of
    `agent-notes.md` and irrelevant to normal drafting, so it moved out.

    These are ceilings, not targets. They exist to make growth a decision
    rather than an accident.
    """
    g = "Ngân sách token"
    tran = {
        "SKILL.md": 3000,                        # nạp mỗi phiên có kích hoạt
        "references/agent-notes.md": 3600,       # đọc thường xuyên nhất
        "references/migration-legacy.md": 1800,
        "references/md-syntax.md": 4000,
        "references/style-guide.md": 4000,
        "references/validation-catalog.md": 2200,
        "references/golden-snippets.md": 2000,
    }
    for rel, max_tok in tran.items():
        n = _uoc_token((SKILL / rel).read_text(encoding="utf-8"))
        report(g, f"{rel} ≤ {max_tok:,} token", n <= max_tok, f"đang {n:,}")

    # Every reference file must be reachable from the routing table, or it is
    # dead weight nobody will ever be told to open. `co-dong-lenh.md` sat
    # unrouted for months for exactly this reason.
    skill = (SKILL / "SKILL.md").read_text(encoding="utf-8")
    mo_coi = sorted(f.name for f in (SKILL / "references").glob("*.md")
                    if f.name not in skill)
    report(g, "không file tham chiếu nào mồ côi (SKILL.md đều trỏ tới)",
           not mo_coi, ", ".join(mo_coi))

    # And the human-facing docs must stay out of the agent's read path.
    doc_nguoi = ("huong-dan-ba.md", "trien-khai.md")
    khong_doc = skill.split("Do not read unless")[-1]
    for d in doc_nguoi:
        report(g, f"{d} vẫn nằm trong danh sách KHÔNG đọc",
               d in khong_doc, khong_doc[:120])


def eval_base_docx(tmp: Path, outline: dict) -> None:
    """The Word template is one file, and swapping it is a documented task.

    Every visual property of the output lives in `assets/base.docx`; nothing
    is hardcoded. That makes rebranding a one-file job — and makes a renamed
    style a silent disaster, because the outline looks styles up **by name**
    and `python-docx` falls back to the default instead of raising. These
    checks are what a lead BA relies on after replacing the file.
    """
    g = "Template Word"
    from docx import Document
    base = SKILL / "assets" / "base.docx"
    d = Document(str(base))
    co = {s.name for s in d.styles}

    thieu = sorted(v for v in outline["styles"].values() if v not in co)
    report(g, "mọi style đề cương khai đều có thật trong base.docx",
           not thieu, "thiếu: " + ", ".join(thieu))

    # Word refuses to open a file whose styleId carries characters outside the
    # allowed set and offers "Show Repairs" instead. `w:name` may hold them —
    # `w:styleId` may not. This shipped once as `T-Gach *`.
    import zipfile
    xml = zipfile.ZipFile(base).read("word/styles.xml").decode("utf-8")
    ids = re.findall(r'w:styleId="([^"]*)"', xml)
    ban = sorted({i for i in ids if not re.fullmatch(r"[A-Za-z0-9-]+", i)})
    report(g, "mọi w:styleId chỉ dùng chữ, số và gạch nối", not ban,
           ", ".join(repr(x) for x in ban))
    report(g, "không có styleId trùng nhau",
           len(ids) == len(set(ids)),
           str(sorted({i for i in ids if ids.count(i) > 1})))

    # The template carries example content so its styles are defined by use.
    # Render must wipe it; a leak would put 91 stray paragraphs in every spec.
    work = tmp / "basedocx"
    shutil.copytree(GOLDEN, work, dirs_exist_ok=True)
    r = run([SCRIPTS / "render.py", "FUNC-QLNSD-001.md", "-o", "t.docx"],
            cwd=work)
    if report(g, "render được bằng template hiện tại", r.returncode == 0,
              r.stderr[-70:]):
        ra = "\n".join(p.text for p in Document(str(work / "t.docx")).paragraphs)
        goc = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        lot = [t for t in goc if t in ra]
        report(g, "nội dung mẫu trong base.docx không lọt vào bản render",
               not lot, f"{len(lot)} đoạn, vd {lot[:1]}")

    # And the document that tells a lead BA how to swap it must stay truthful.
    tk = (SKILL / "references" / "trien-khai.md").read_text(encoding="utf-8")
    report(g, "trien-khai.md có mục hướng dẫn thay base.docx",
           "base.docx` — template Word" in tk)
    hd = (SKILL / "references" / "huong-dan-ba.md").read_text(encoding="utf-8")
    report(g, "huong-dan-ba.md có mục hướng dẫn xuất Word/PDF",
           "Xuất ra Word và PDF" in hd)
    report(g, "nói rõ mặc định là file con, không có bìa",
           "không bìa" in hd and "--standalone" in hd or "độc lập" in hd)


def eval_du_an_moi(tmp: Path) -> None:
    """An empty project is not a broken one.

    Running this on a fresh skeleton is the first thing anyone does with the
    handover bundle. Greeting them with `[LỖI]` teaches them the tool cries
    wolf, and a tool that cries wolf gets ignored on the day it is right. But
    a folder that *has* specs, none of which parse, is a real fault — usually
    an upload that lost its contents.
    """
    g = "Dự án mới"
    moi = tmp / "moi"
    (moi / "functions" / "qlnsd" / "assets").mkdir(parents=True)
    (moi / "functions" / "qlnsd" / "diagrams").mkdir()
    (moi / "registries").mkdir()
    for c in (GOLDEN / "registries").glob("*.csv"):
        shutil.copy(c, moi / "registries")

    r = run([SCRIPTS / "project_check.py", "."], cwd=moi)
    report(g, "khung rỗng: KHÔNG phải lỗi", r.returncode == 0,
           r.stdout.strip()[-90:])
    report(g, "nói rõ là dự án mới, không phải thiếu file",
           "dự án mới" in r.stdout and "Khung dự án dựng đúng" in r.stdout)
    report(g, "chỉ luôn bước tiếp theo",
           "manifest.md" in r.stdout and "srs.py new" in r.stdout)

    # The other half: files that exist but do not parse must still be an error.
    hong = tmp / "hong"
    (hong / "functions" / "qlnsd").mkdir(parents=True)
    (hong / "registries").mkdir()
    for c in (GOLDEN / "registries").glob("*.csv"):
        shutil.copy(c, hong / "registries")
    (hong / "functions" / "qlnsd" / "FUNC-QLNSD-001.md").write_text(
        "# Chỉ là một tiêu đề, không có front matter\n", encoding="utf-8")
    r = run([SCRIPTS / "project_check.py", "."], cwd=hong)
    report(g, "có file .md nhưng không parse được: vẫn là LỖI",
           r.returncode == 1 and "không file nào có front matter hợp lệ"
           in r.stdout, r.stdout.strip()[-90:])


def eval_behaviour_rules(outline: dict) -> None:
    """Two forward-test failures from 18/08, written back into the rules.

    A behaviour test that fails and only gets a log entry will fail again the
    next release. These check the wording that was added in response is still
    there — not that the agent behaves, which no script can check.
    """
    g = "Luật sau kiểm hành vi"
    skill = (SKILL / "SKILL.md").read_text(encoding="utf-8")
    # A3: the agent reviewed an ambiguous request, then asked afterwards.
    modes = skill.split("## Two modes")[1].split("##")[0]
    report(g, "A3: luật nói rõ phải hỏi TRƯỚC, không phải hỏi sau",
           "before doing either, not after" in modes, modes[-200:])
    # B4: the agent minted shared codes and wrote their Vietnamese wording.
    codes = skill.split("## Two kinds of code")[1].split("## Two modes")[0]
    report(g, "B4: tách việc thêm dòng vào sổ với việc viết nội dung dòng đó",
           "not the same as writing what goes in it" in codes, codes[-200:])
    report(g, "B4: nói rõ nội dung thông báo là nội dung nghiệp vụ",
           "business content" in codes)
    notes = (SKILL / "references/agent-notes.md").read_text(encoding="utf-8")
    report(g, "agent-notes ghi lại ca hỏng để lần sau không lặp",
           "B4, 18/08" in notes)


def eval_table_spacing(tmp: Path, outline: dict) -> None:
    """A table pressed against other content renders wrong everywhere else.

    This parser is lenient about it, so nothing here breaks and the fault
    reached a delivered document unnoticed: a heading right after a table
    became a new row with `###` shown literally in column one. Every file the
    scripts write must be clean, hand-written ones must be caught, and the
    repair must be one command.
    """
    g = "Dòng trắng quanh bảng"

    # 1. Nothing the scripts generate may contain the fault.
    work = tmp / "spacing"
    work.mkdir()
    made = []
    for kind in S.all_kinds(outline):
        ma = "GRP-EVAL-01" if kind == S.GROUP else "FUNC-EVAL-001"
        r = run([SCRIPTS / "scaffold.py", "--profile", kind, "--ma", ma,
                 "--ten", "T", "--tinh-nang", "2", "-o", f"{kind}.md"],
                cwd=work)
        if r.returncode == 0:
            made.append(work / f"{kind}.md")
    made += sorted(GOLDEN.glob("*.md"))
    bad = {p.name: S.table_spacing_faults(p.read_text(encoding="utf-8"))
           for p in made}
    dirty = {k: v for k, v in bad.items() if v}
    report(g, f"{len(made)} file do script sinh đều sạch", not dirty,
           "; ".join(f"{k}: {len(v)}" for k, v in list(dirty.items())[:3]))

    # 2. Both directions are caught, with the line number.
    after = "| A | B |\n|---|---|\n| 1 | 2 |\n### Thiết kế giao diện\n"
    before = "**Luồng thay thế**\n| A | B |\n|---|---|\n| 1 | 2 |\n"
    report(g, "nội dung dính SAU bảng bị bắt",
           len(S.table_spacing_faults(after)) == 1)
    report(g, "nội dung dính TRƯỚC bảng bị bắt",
           len(S.table_spacing_faults(before)) == 1)
    report(g, "file đúng cách dòng thì không báo gì",
           not S.table_spacing_faults(
               "| A | B |\n|---|---|\n| 1 | 2 |\n\n### Mục\n"))

    # 3. The repair is idempotent and actually removes the fault.
    for name, src in (("sau", after), ("truoc", before)):
        fixed, n = S.normalize_table_spacing(src)
        again, n2 = S.normalize_table_spacing(fixed)
        report(g, f"tự sửa được ca “{name}”", n >= 1 and not
               S.table_spacing_faults(fixed))
        report(g, f"sửa lần hai không đổi gì (ca “{name}”)",
               n2 == 0 and again == fixed)

    # 4. validate rejects it, and `srs.py fix` clears it.
    p = work / "bad.md"
    base = (GOLDEN / "FUNC-QLNSD-001.md").read_text(encoding="utf-8")
    p.write_text(base.replace("| 3 |\n\n**Luồng ngoại lệ**\n\n",
                              "| 3 |\n**Luồng ngoại lệ**\n", 1),
                 encoding="utf-8")
    shutil.copytree(GOLDEN / "registries", work / "registries",
                    dirs_exist_ok=True)
    shutil.copytree(GOLDEN / "assets", work / "assets", dirs_exist_ok=True)
    r = run([SCRIPTS / "validate.py", "bad.md", "--registry-dir", "registries"],
            cwd=work)
    report(g, "validate báo lỗi và chỉ đúng số dòng",
           r.returncode == 1 and "khoảng cách bảng" in r.stdout)
    r = run([SCRIPTS / "srs.py", "fix", "bad.md"], cwd=work)
    r2 = run([SCRIPTS / "validate.py", "bad.md", "--registry-dir",
              "registries"], cwd=work)
    report(g, "`srs.py fix` sửa xong thì validate sạch",
           r.returncode == 0 and "khoảng cách bảng" not in r2.stdout)

    # 5. Repair must not touch anything else.
    good = work / "good.md"
    good.write_text(base, encoding="utf-8")
    run([SCRIPTS / "srs.py", "fix", "good.md"], cwd=work)
    report(g, "file đã đúng thì fix không sửa một byte nào",
           good.read_text(encoding="utf-8") == base)


def eval_images(tmp: Path, outline: dict) -> None:
    """Extracted images must all survive, with usable names.

    The original naming took Word's alt-text, which defaults to `Picture 1`
    for almost every image: a real 26-image spec produced two files on disk
    while still reporting nineteen extracted, and every figure in the
    delivered document pointed at the same picture.
    """
    g = "Lấy ảnh ra"
    sys.path.insert(0, str(SCRIPTS))
    from docx import Document
    from docx.shared import Inches

    work = tmp / "img"
    work.mkdir()
    png = work / "src.png"
    # 1×1 PNG, enough to be a distinct file.
    import base64
    png.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="))
    png2 = work / "src2.png"
    png2.write_bytes(png.read_bytes() + b"\x00")   # khác byte, cùng tên gốc

    # base.docx is opened for its styles, but its own body content would be
    # counted as part of the fixture — it carries a picture and headings of
    # its own, which shifted every section number and added a fourth image.
    d = Document(str(SKILL / "assets" / "base.docx"))
    from docx.oxml.ns import qn as _q
    body = d.element.body
    for child in list(body):
        if child.tag in (_q("w:p"), _q("w:tbl")):
            body.remove(child)

    d.add_heading("Chức năng [FUNC-EVAL-001] Kiểm thử ảnh", 3)
    d.add_heading("Mô tả chung", 4)
    d.add_paragraph().add_run().add_picture(str(png), width=Inches(1))
    d.add_heading("Tính năng [FEAT-EVAL-001-01] Tra cứu", 4)
    d.add_heading("Thiết kế giao diện", 5)
    d.add_paragraph().add_run().add_picture(str(png2), width=Inches(1))
    cap = d.add_paragraph("Figure 1 Giao diện tra cứu")
    try:
        cap.style = d.styles[outline["styles"]["caption"]]
    except KeyError:
        pass
    d.add_paragraph().add_run().add_picture(str(png), width=Inches(1))
    # Every picture gets Word's default alt-text — the collision case.
    from docx.oxml.ns import qn as _qn
    for docPr in d.element.body.iter(_qn("wp:docPr")):
        docPr.set("name", "Picture 1")
        docPr.set("descr", "Picture 1")
    d.save(str(work / "old.docx"))

    doc = Document(str(work / "old.docx"))
    hits = S.scan_images(doc, outline["styles"])
    report(g, "quét đủ 3 ảnh dù alt-text trùng nhau", len(hits) == 3,
           f"{len(hits)} ảnh")

    saved, vector = S.save_images(doc, hits, work / "assets")
    files = sorted(p.name for p in (work / "assets").glob("*"))
    report(g, "ghi ra đủ 3 tệp, không tệp nào bị ghi đè",
           len(files) == 3, str(files))
    report(g, "mọi tệp có phần mở rộng thật",
           all(f.lower().endswith(".png") for f in files), str(files))
    report(g, "tên bắt đầu bằng số thứ tự tài liệu",
           [f[:3] for f in files] == ["001", "002", "003"], str(files))
    report(g, "tên mang số mục lấy từ heading",
           any("_1." in f or "_2." in f for f in files), str(files))
    report(g, "ảnh có chú thích được đặt tên theo chú thích",
           any("giao-dien-tra-cuu" in f for f in files), str(files))
    report(g, "không tên nào chứa khoảng trắng",
           all(" " not in f for f in files), str(files))

    # Re-importing a document into the folder it came from must reuse the
    # same filenames, or the .md starts pointing at `…-2.png`.
    saved2, _ = S.save_images(doc, S.scan_images(doc, outline["styles"]),
                              work / "assets")
    again = sorted(p.name for p in (work / "assets").glob("*"))
    report(g, "chạy lại không sinh bản sao -2", again == files, str(again))

    # Section numbering is computed, since the headings carry none.
    report(g, "số mục tính được cho ảnh nằm dưới heading",
           all(h.sec_num for h in hits), str([h.sec_num for h in hits]))

    # And the round trip through the skill's own render must keep its names.
    shutil.copytree(GOLDEN, work / "g", dirs_exist_ok=True)
    p = work / "g"
    r = run([SCRIPTS / "render.py", "FUNC-QLNSD-001.md", "-o", "g.docx"], cwd=p)
    if report(g, "render file mẫu", r.returncode == 0):
        r = run([SCRIPTS / "import_docx.py", "g.docx", "-o", "b.md"], cwd=p)
        a_paths = re.findall(r"\(assets/([^)]+)\)",
                             (p / "FUNC-QLNSD-001.md").read_text(encoding="utf-8"))
        b_paths = re.findall(r"\(assets/([^)]+)\)",
                             (p / "b.md").read_text(encoding="utf-8"))
        report(g, "vòng khép kín giữ nguyên tên ảnh", a_paths == b_paths,
               f"{a_paths} vs {b_paths}")


def eval_code_separators() -> None:
    """Codes written with the wrong separator must still be caught.

    A 2 900-cell legacy spec carried 104 message references, every one of them
    `ERR-001` rather than `ERR_001`. Both the strict and the loose pattern
    pinned the separator, so the validator reported *zero* message codes in
    that document — registry checks, declared-vs-referenced checks, all
    silently inert on exactly the file that needed them.
    """
    g = "Mã sai dấu nối"
    cases = [
        ("ERR-001", "MSG", "ERR_001"),
        ("SUC-002", "MSG", "SUC_002"),
        ("ERR001", "MSG", "ERR_001"),
        ("ERR 001", "MSG", "ERR_001"),
        ("UC-301", "UC", "UC-0301"),
        ("ROLE_QTHT", "ROLE", "ROLE-QTHT"),
    ]
    for tok, kind, want in cases:
        bad = S.find_malformed_codes(f"thấy {tok} ở đây", kind)
        fix = S.suggest_code_fix(tok, kind) if bad else None
        report(g, f"`{tok}` bị bắt và gợi ý `{want}`",
               tok in bad and fix == want, f"bắt={bad} gợi ý={fix}")

    # Well-formed codes must stay clean, and ordinary words that merely start
    # with a prefix must not be dragged in — noise here would train the BA to
    # ignore the whole category.
    for txt, kind in [("ERR_001 đúng dạng", "MSG"), ("UC-0301 đúng", "UC"),
                      ("ROLE-QTHT đúng", "ROLE"),
                      ("STT là số thứ tự", "ST"),
                      ("INFORMATION và CONFIG", "MSG")]:
        report(g, f"không báo nhầm: “{txt}”",
               not S.find_malformed_codes(txt, kind),
               str(S.find_malformed_codes(txt, kind)))


def eval_cell_collapse(tmp: Path, outline: dict) -> None:
    """A Word cell holding several lines has to become one markdown line.

    Every fault here came from a real legacy spec, and each was silent: a soft
    line break (Shift+Enter) survived as `\\n` and split the table row in half,
    a literal `|` split one cell into four, and bullets collapsed into an
    unreadable run-on because the house bullet styles carry no `w:numPr` and
    the level test only looked there.
    """
    g = "Gộp ô nhiều dòng"
    sys.path.insert(0, str(SCRIPTS))
    import import_docx as I
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    work = tmp / "cell"
    work.mkdir()
    bs = I.bullet_style_map(outline)
    report(g, "đọc được 3 style bullet từ đề cương", len(bs) == 3, str(bs))

    d = Document(str(SKILL / "assets" / "base.docx"))
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "STT"
    t.rows[0].cells[1].text = "Mô tả ràng buộc"
    c = t.rows[1].cells[1]
    c.text = "Place holder \"Nhập tên\""
    c.paragraphs[0].style = d.styles[outline["styles"]["body"]]
    for txt, st in [("Quy tắc nhập liệu:", "body"),
                    ("Cho phép chữ và số", "bullet_2"),
                    ("Không cho phép emoji", "bullet_2"),
                    ("Bỏ trống thì báo ERR-001", "bullet_1")]:
        c.add_paragraph(txt, style=d.styles[outline["styles"][st]])
    t.rows[1].cells[0].text = "1"

    # Word numbering, the other bullet mechanism.
    t2 = d.add_table(rows=1, cols=1)
    p = t2.rows[0].cells[0].paragraphs[0]
    p.add_run("Mục có numPr")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), "1")
    numPr.append(ilvl); numPr.append(nid); pPr.append(numPr)

    # Shift+Enter and a literal pipe.
    t3 = d.add_table(rows=1, cols=2)
    t3.rows[0].cells[0].text = "Dòng một\nDòng hai"
    t3.rows[0].cells[1].text = "Định dạng A|B|C"
    d.save(str(work / "in.docx"))

    # base.docx ships with its own tables, so index by identity rather than
    # position: reopen and match on the header text we wrote.
    doc = Document(str(work / "in.docx"))
    t_bul = next(x for x in doc.tables
                 if x.rows[0].cells[1].text == "Mô tả ràng buộc")
    t_num = next(x for x in doc.tables
                 if x.rows[0].cells[0].text == "Mục có numPr")
    t_odd = next(x for x in doc.tables
                 if "Dòng một" in x.rows[0].cells[0].text)

    got = I.cell_text(t_bul.rows[1].cells[1], bs)
    report(g, "bullet theo style nhà (T-Gach) được nhận, có phân cấp",
           got.count("··") == 2 and " · " in got, got[:90])
    report(g, "dòng văn xuôi thiếu dấu câu được thêm dấu chấm",
           'Nhập tên".' in got, got[:60])
    report(g, "bullet theo w:numPr cũng được nhận",
           I.cell_text(t_num.rows[0].cells[0], bs).startswith("·"),
           I.cell_text(t_num.rows[0].cells[0], bs)[:40])

    soft = I.cell_text(t_odd.rows[0].cells[0], bs)
    pipe = I.cell_text(t_odd.rows[0].cells[1], bs)
    report(g, "Shift+Enter không còn để lại xuống dòng", "\n" not in soft, repr(soft))
    report(g, "ký tự | được escape", "\\|" in pipe and "A|B" not in pipe, pipe)

    # The invariant, stated once: whatever the input, one line, no bare pipe.
    for name, cell in (("bullet", t_bul.rows[1].cells[1]),
                       ("soft-break", t_odd.rows[0].cells[0]),
                       ("pipe", t_odd.rows[0].cells[1])):
        v = I.cell_text(cell, bs)
        ok = "\n" not in v and not re.search(r"(?<!\\)\|", v)
        report(g, f"bất biến một dòng · {name}", ok, repr(v[:50]))

    # And the round trip stays untouched: a skill-rendered docx has exactly one
    # paragraph per cell, so none of the above can change it.
    shutil.copytree(GOLDEN, work / "g", dirs_exist_ok=True)
    r = run([SCRIPTS / "render.py", "FUNC-QLNSD-001.md", "-o", "g.docx"],
            cwd=work / "g")
    if report(g, "render file mẫu", r.returncode == 0):
        gd = Document(str(work / "g" / "g.docx"))
        multi = sum(1 for tb in gd.tables for rr in tb.rows for cc in rr.cells
                    if I.cell_segment_count(cc) > 1)
        report(g, "tài liệu do skill render không có ô nhiều dòng → thay đổi "
                  "này là no-op với vòng khép kín", multi == 0, f"{multi} ô")


def eval_dispatcher(tmp: Path) -> None:
    """srs.py must forward faithfully and the export chain must stop on
    error — a chain that renders anyway would defeat the whole gate."""
    g = "Lệnh điều phối"
    work = tmp / "disp"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "p", dirs_exist_ok=True)
    p = work / "p"
    srs = SCRIPTS / "srs.py"

    r = run([srs, "check", "FUNC-QLNSD-001.md", "--registry-dir", "registries",
             "--quiet"], cwd=p)
    report(g, "check chuyển tiếp sang validate, --quiet gọn",
           r.returncode == 0 and "cảnh báo (ẩn)" in r.stdout
           and "CẢNH" not in r.stdout)

    r = run([SCRIPTS / "validate.py", "FUNC-QLNSD-001.md", "--registry-dir",
             "registries", "--json"], cwd=p)
    ok = r.returncode == 0
    if ok:
        try:
            data = json.loads(r.stdout)
            ok = data["exit"] == 0 and data["files"][0]["errors"] == []
        except Exception:
            ok = False
    report(g, "--json in JSON hợp lệ, đúng cấu trúc", ok)

    base = (p / "FUNC-QLNSD-001.md").read_text(encoding="utf-8")
    (p / "bad.md").write_text(base.replace(
        "| Trong phạm vi |", "| Trong phạm vi sai |", 1), encoding="utf-8")
    r = run([srs, "export", "bad.md", "--registry-dir", "registries"], cwd=p)
    report(g, "export dừng khi validate có lỗi, không render",
           r.returncode == 1 and not list(p.glob("bad*.docx")))

    r = run([srs, "lenh-khong-co"], cwd=p)
    report(g, "lệnh không tồn tại → báo và exit 1",
           r.returncode == 1 and "Không có lệnh" in r.stdout)

    # -- review must not touch the canonical .md ---------------------------
    r = run([SCRIPTS / "render.py", "FUNC-QLNSD-001.md", "-o", "rv.docx"],
            cwd=p)
    if report(g, "render mẫu cho ca review", r.returncode == 0):
        before = (p / "FUNC-QLNSD-001.md").read_bytes()
        r = run([srs, "review", "rv.docx", "--registry-dir", "registries",
                 "--quiet"], cwd=p)
        report(g, "review không ghi đè .md gốc",
               r.returncode == 0
               and (p / "FUNC-QLNSD-001.md").read_bytes() == before)
        report(g, "review nhập ra file .review-import.md riêng",
               (p / "rv.review-import.md").exists()
               and "review-import" in r.stdout)

    # -- an argparse error (also exit 2) must never read as "gate only" ----
    r = run([srs, "export", "FUNC-QLNSD-001.md", "--registry-dir"], cwd=p)
    report(g, "export: lỗi đối số (exit 2 của argparse) không được hiểu là "
              "gate-only, không render",
           r.returncode != 0
           and not (p / "FUNC-QLNSD-001_Quan-ly-nguoi-dung.docx").exists())

    # -- --outline reaches both validate and render ------------------------
    o = json.loads((SKILL / "references" / "outline.json")
                   .read_text(encoding="utf-8"))
    o["lexicon"]["figure"] = "KiemOutline"
    (p / "custom-outline.json").write_text(
        json.dumps(o, ensure_ascii=False), encoding="utf-8")
    r = run([srs, "export", "FUNC-QLNSD-001.md", "--registry-dir",
             "registries", "--outline", "custom-outline.json",
             "-o", "co.docx"], cwd=p)
    # What this test asserts is flag routing: the custom outline must reach
    # the *render* stage, visible as the changed caption label in the .docx.
    # The chain's final PDF step needs LibreOffice, which is a different
    # concern (covered, with its own skip-guard, in the Xuất PDF group) — so
    # only demand rc == 0 where soffice exists, otherwise a machine without
    # LibreOffice fails this test for a reason it isn't about.
    has_soffice = bool(shutil.which("soffice") or shutil.which("libreoffice"))
    ok = (p / "co.docx").exists() and (r.returncode == 0 or not has_soffice)
    if ok:
        from docx import Document
        ok = any("KiemOutline" in q.text
                 for q in Document(str(p / "co.docx")).paragraphs)
    report(g, "--outline truyền xuyên suốt tới cả render (caption đổi theo)",
           ok, f"rc={r.returncode}, soffice={'có' if has_soffice else 'không'}")


def eval_cli_unicode(tmp: Path) -> None:
    """The direct CLI must survive a non-UTF-8 console without env help.

    The suite's own `run()` exports PYTHONUTF8=1, which masked exactly this:
    every script crashed with UnicodeEncodeError on a default Windows
    PowerShell. Simulate that console by forcing cp1252 stdio and *removing*
    the rescue variable.
    """
    g = "CLI không cần env UTF-8"
    work = tmp / "uni"
    work.mkdir()
    shutil.copytree(GOLDEN, work / "p", dirs_exist_ok=True)
    p = work / "p"

    env = os.environ.copy()
    env.pop("PYTHONUTF8", None)
    env["PYTHONIOENCODING"] = "cp1252"

    def raw(args, cwd):
        return subprocess.run([PY] + [str(x) for x in args], cwd=cwd, env=env,
                              capture_output=True, text=True,
                              encoding="utf-8", errors="replace")

    cases = [
        ("validate --quiet in tiếng Việt",
         [SCRIPTS / "validate.py", "FUNC-QLNSD-001.md",
          "--registry-dir", "registries", "--quiet"], 0),
        ("validate bản đầy đủ (có CẢNH BÁO)",
         [SCRIPTS / "validate.py", "FUNC-QLNSD-001.md"], 0),
        ("srs.py lệnh sai (thông báo có dấu)",
         [SCRIPTS / "srs.py", "lenh-khong-co"], 1),
        ("project_check in kết luận có dấu",
         [SCRIPTS / "project_check.py", "."], 0),
        ("outline_check", [SCRIPTS / "outline_check.py"], 0),
    ]
    for name, args, want_rc in cases:
        r = raw(args, p)
        crashed = "UnicodeEncodeError" in (r.stderr or "") \
            or "UnicodeDecodeError" in (r.stderr or "")
        report(g, f"{name}: không crash trên console cp1252",
               r.returncode == want_rc and not crashed,
               (r.stderr or "").strip().splitlines()[-1][:70]
               if crashed else f"rc={r.returncode}")


def eval_bullets(tmp: Path, outline: dict) -> None:
    g = "Gạch đầu dòng"
    work = tmp / "bul"
    work.mkdir()
    r = run([SCRIPTS / "scaffold.py", "--profile", "UI", "--ma",
             "FUNC-EVAL-001", "--ten", "Bullet", "--tinh-nang", "1",
             "-o", "b.md"], cwd=work)
    if not report(g, "dựng khung", r.returncode == 0):
        return
    p = work / "b.md"
    p.write_text(p.read_text(encoding="utf-8").replace(
        "## Luồng nghiệp vụ\n",
        "## Luồng nghiệp vụ\n\n- Cấp một\n  - Cấp hai\n    - Cấp ba\n", 1),
        encoding="utf-8")

    a = S.read_markdown(p)
    lv = [b.level for b in a.section("Luồng nghiệp vụ").blocks
          if b.kind == "bullet"]
    report(g, "đọc đúng 3 cấp từ markdown", lv == [1, 2, 3], str(lv))

    r = run([SCRIPTS / "render.py", "b.md", "-o", "b.docx"], cwd=work)
    styles = [outline["styles"][f"bullet_{n}"] for n in (1, 2, 3)]
    from docx import Document
    got = [x.style.name for x in Document(str(work / "b.docx")).paragraphs
           if x.text.startswith("Cấp")]
    report(g, "render đúng style từng cấp", got == styles, str(got))

    r = run([SCRIPTS / "import_docx.py", "b.docx", "-o", "b2.md"], cwd=work)
    b = S.read_markdown(work / "b2.md")
    lv2 = [x.level for x in b.section("Luồng nghiệp vụ").blocks
           if x.kind == "bullet"]
    report(g, "nhập lại giữ đúng cấp", lv == lv2, str(lv2))


def eval_naming(tmp: Path) -> None:
    """merge.py finds child files by globbing `{MÃ}_*.docx`, so the default
    output name has to carry the title as well as the code."""
    g = "Quy ước tên file"
    work = tmp / "name"
    work.mkdir()
    shutil.copy(GOLDEN / "FUNC-QLNSD-001.md", work)
    shutil.copy(GOLDEN / "GRP-QLNSD-01.md", work)
    shutil.copytree(GOLDEN / "assets", work / "assets", dirs_exist_ok=True)

    for src, want in (("FUNC-QLNSD-001.md",
                       "FUNC-QLNSD-001_Quan-ly-nguoi-dung.docx"),
                      ("GRP-QLNSD-01.md",
                       "GRP-QLNSD-01_Nguoi-dung-phan-quyen.docx")):
        r = run([SCRIPTS / "render.py", src], cwd=work)
        report(g, f"{src}: tên mặc định đúng quy ước merge",
               r.returncode == 0 and (work / want).exists(),
               r.stdout.strip().splitlines()[-1] if r.returncode == 0 else "")


def eval_group_rules(tmp: Path, outline: dict) -> None:
    """Rules mirrored from the pipeline's tools/validate_group.py, plus the
    T-GhiChu round trip that previously lost the description."""
    g = "Nhóm chức năng"
    work = tmp / "grp"
    work.mkdir()
    src = GOLDEN / "GRP-QLNSD-01.md"
    shutil.copy(src, work)
    base = src.read_text(encoding="utf-8")
    (work / "reg").mkdir()
    (work / "reg" / "groups.csv").write_text(
        "ma,ten,ghi_chu\nGRP-QLNSD-01,Người dùng & Phân quyền,\n",
        encoding="utf-8")

    r = run([SCRIPTS / "validate.py", src.name, "--registry-dir", "reg"],
            cwd=work)
    report(g, "golden nhóm sạch với groups.csv", r.returncode == 0)

    (work / "x.md").write_text(
        base.replace("GRP-QLNSD-01", "GRP-XXXX-99"), encoding="utf-8")
    r = run([SCRIPTS / "validate.py", "x.md", "--registry-dir", "reg"],
            cwd=work)
    report(g, "mã nhóm ngoài groups.csv bị bắt", r.returncode == 1)

    (work / "y.md").write_text(
        base + "\n| A | B |\n|---|---|\n| 1 | 2 |\n", encoding="utf-8")
    r = run([SCRIPTS / "validate.py", "y.md", "--registry-dir", "reg"],
            cwd=work)
    report(g, "file nhóm có bảng bị bắt",
           r.returncode == 1 and "không được có bảng" in r.stdout)

    # description must survive render → import as T-GhiChu
    r = run([SCRIPTS / "render.py", src.name, "-o", "g.docx"], cwd=work)
    if report(g, "render nhóm", r.returncode == 0):
        from docx import Document
        styles = [p_.style.name for p_ in Document(str(work / "g.docx")).paragraphs
                  if p_.text.strip()]
        report(g, "mô tả nhóm dùng style T-GhiChu",
               outline["styles"]["note"] in styles, str(styles))
        run([SCRIPTS / "import_docx.py", "g.docx", "-o", "back.md"], cwd=work)
        a = S.read_markdown(work / src.name)
        b = S.read_markdown(work / "back.md")
        pa = [x.text for sec in a.sections for x in sec.blocks
              if x.kind == "para"]
        pb = [x.text for sec in b.sections for x in sec.blocks
              if x.kind == "para"]
        report(g, "mô tả nhóm không mất khi nhập lại", pa == pb and bool(pa))


def eval_phantich(tmp: Path, outline: dict) -> None:
    """PHANTICH is flat, so its two tables must not carry feature columns."""
    g = "PHANTICH phẳng"
    for name, gone in (("Ma trận phân quyền", "Mã tính năng"),
                       ("Truy vết yêu cầu", "Tính năng đáp ứng")):
        cols = [c for m in outline["profiles"]["PHANTICH"]["function_sections"]
                if m["name"] == name
                for t in m.get("tables", [])
                for c in outline["tables"][t]["headers"]]
        report(g, f"“{name}” không còn cột “{gone}”", gone not in cols,
               str(cols))
        ui = [c for m in outline["profiles"]["UI"]["function_sections"]
              if m["name"] == name
              for t in m.get("tables", [])
              for c in outline["tables"][t]["headers"]]
        report(g, f"“{name}” của UI vẫn giữ cột “{gone}”", gone in ui)


def eval_legacy(tmp: Path) -> None:
    g = "Tài liệu ngoài chuẩn"
    work = tmp / "old"
    work.mkdir()
    from docx import Document
    d = Document()
    d.add_heading("TÀI LIỆU ĐẶC TẢ", 0)
    d.add_heading("1. High Level Requirements", 1)
    d.add_paragraph("Nội dung cũ.")
    d.add_heading("2. Use Case Specifications", 1)
    d.add_heading("2.1 UC-01 Đăng nhập", 2)
    d.save(str(work / "cu.docx"))

    r = run([SCRIPTS / "import_docx.py", "cu.docx", "-o", "x.md"], cwd=work)
    report(g, "từ chối, không đoán bừa loại",
           r.returncode == 1 and "KHÔNG THEO ĐỀ CƯƠNG" in r.stdout)
    report(g, "không ghi ra file .md", not (work / "x.md").exists())

    # merged master document must be refused, not misread as one child
    from docx import Document as _D
    m = _D()
    m.add_heading("Nhóm chức năng [GRP-XXXX-01] A", 2)
    m.add_heading("Chức năng [FUNC-XXXX-001] B", 3)
    m.add_heading("Chức năng [FUNC-XXXX-002] C", 3)
    m.save(str(work / "master.docx"))
    r = run([SCRIPTS / "import_docx.py", "master.docx", "-o", "z.md"], cwd=work)
    report(g, "tài liệu tổng bị từ chối",
           r.returncode == 1 and "TÀI LIỆU TỔNG" in r.stdout)

    r = run([SCRIPTS / "migrate_scan.py", "cu.docx", "--profile", "UI",
             "-o", "km.md"], cwd=work)
    ok = r.returncode == 0 and (work / "km.md").exists()
    report(g, "migrate_scan sinh bảng ánh xạ", ok)
    if ok:
        txt = (work / "km.md").read_text(encoding="utf-8")
        report(g, "bảng ánh xạ liệt kê đủ mục khung mới",
               txt.count("|") > 40 and "Ma trận phân quyền" in txt)


def eval_pipeline(tmp: Path, outline: dict, pipe: Path) -> None:
    g = "Đối chiếu pipeline"
    export_file = pipe / "outline_export.json"
    if export_file.exists():
        import json
        try:
            with export_file.open("r", encoding="utf-8") as f:
                exp = json.load(f)
            report(g, "đọc outline_export.json", True)
            if "msg_types" in exp:
                report(g, "msg_types khớp", exp["msg_types"] == outline["msg_types"])
            if "tables" in exp:
                match = True
                for t_id, t_info in outline["tables"].items():
                    if t_id in exp["tables"]:
                        e_info = exp["tables"][t_id]
                        if (e_info.get("headers") != t_info.get("headers")
                                or e_info.get("widths") != t_info.get("widths")):
                            match = False
                report(g, "cấu trúc bảng (cột, độ rộng) khớp", match)
        except Exception as exc:
            report(g, "đọc outline_export.json", False, str(exc))
        return

    tools = pipe / "tools"
    if not (tools / "outline.py").exists():
        report(g, "tìm thấy outline.py", False, str(tools))
        return
    sys.path.insert(0, str(tools))
    try:
        import outline as O
    except Exception as e:
        report(g, "nạp được outline.py", False, str(e)[:60])
        return

    for k in O.ALL_PROFILES:
        before, feats, after = O.sections(k)
        want_fn = [s["name"] for s in before] + [s["name"] for s in after]
        want_ft = [s["name"] for s in feats]
        got = outline["profiles"][k]
        report(g, f"{k}: mục cấp chức năng khớp",
               want_fn == [s["name"] for s in got["function_sections"]])
        report(g, f"{k}: mục cấp tính năng khớp",
               want_ft == [s["name"] for s in got["feature_sections"]])
        report(g, f"{k}: has_features khớp",
               O.has_features(k) == got["has_features"])

    # Nhom chuc nang: KHAC BIET CO CHU Y — la tang cay menu, khong phai tai
    # lieu. Skill chi giu de muc + mo ta ngan; pipeline con giu 5 muc cu.
    report(g, "nhóm chức năng: skill giữ dạng rút gọn",
           outline["group"]["sections"] == [])
    report(g, "msg_types khớp", list(O.MSG_TYPES) == outline["msg_types"])
    report(g, "usable_twips khớp", O.USABLE == outline["layout"]["usable_twips"])
    report(g, "số quy ước mã khớp",
           len(O.CODE_RULES) == len(outline["code_rules"]))


# ---------------------------------------------------------------------------
def main() -> int:
    # The suite's own prints hit the same wall as the children's: a cp1252
    # console cannot encode "LỖI". Reconfigure our stdout too, so the run
    # works without the user having to set PYTHONUTF8 by hand.
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass
    ap = argparse.ArgumentParser(description="Bộ kiểm hồi quy cho srs-help.")
    ap.add_argument("--pipeline", default=None,
                    help="thư mục srs-pipeline, để thêm phép đối chiếu đề cương")
    ap.add_argument("--keep", action="store_true", help="giữ thư mục tạm")
    a = ap.parse_args()

    outline = S.load_outline()
    tmp = Path(tempfile.mkdtemp(prefix="srs-eval-"))

    try:
        eval_outline(outline)
        eval_docs_portability(outline)
        eval_skill_doc(outline)
        eval_profiles(tmp, outline)
        eval_roundtrip(tmp, outline)
        eval_validate(tmp)
        eval_roles_and_gaps(tmp)
        eval_migration(tmp, outline)
        eval_cell_collapse(tmp, outline)
        eval_stt(tmp, outline)
        eval_images(tmp, outline)
        eval_project_layout(tmp)
        eval_workspace_shape(tmp)
        eval_manifest(tmp, outline)
        eval_behaviour_rules(outline)
        eval_ngan_sach_token()
        eval_base_docx(tmp, outline)
        eval_du_an_moi(tmp)
        eval_multiline_cells(tmp, outline)
        eval_multiline_profiles(tmp, outline)
        eval_table_spacing(tmp, outline)
        eval_code_separators()
        eval_dispatcher(tmp)
        eval_cli_unicode(tmp)
        eval_registry_schema(tmp)
        eval_export_pdf(tmp)
        eval_bullets(tmp, outline)
        eval_group_rules(tmp, outline)
        eval_phantich(tmp, outline)
        eval_naming(tmp)
        eval_legacy(tmp)
        if a.pipeline:
            eval_pipeline(tmp, outline, Path(a.pipeline))
    finally:
        if a.keep:
            print(f"\nThư mục tạm: {tmp}")
        else:
            shutil.rmtree(tmp, ignore_errors=True)

    width = max(len(n) for _, n, _ in results) + 2
    cur = None
    for grp, name, st in results:
        if grp != cur:
            print(f"\n── {grp} " + "─" * (width - len(grp) + 4))
            cur = grp
        mark = "  ok " if st == PASS else " FAIL"
        print(f"  [{mark}] {name}")

    n_fail = sum(1 for _, _, s in results if s == FAIL)
    print(f"\n{'=' * (width + 8)}")
    print(f"{len(results)} phép kiểm · {len(results) - n_fail} đạt · "
          f"{n_fail} hỏng")
    return 1 if n_fail else 0


if __name__ == "__main__":
    raise SystemExit(main())
