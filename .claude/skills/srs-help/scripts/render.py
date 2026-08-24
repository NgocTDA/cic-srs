#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render.py — .md  ->  .docx theo style cua assets/base.docx.

Nguyen tac: tieu de muc, so cot bang, nhan caption deu lay tu outline.json.
File .md chi cung cap NOI DUNG. Nho vay tai lieu khong the lech chuan du BA go
sai tieu de trong .md — validate.py chan truoc, render.py do lai tu de cuong.

    python render.py FUNC-QLNSD-001.md -o out/FUNC-QLNSD-001.docx
    python render.py FUNC-QLNSD-001.md --standalone --draft
"""
from __future__ import annotations

import argparse
import hashlib
import os
import re
import shutil
import subprocess
import sys
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, Twips

import srslib as S

EMU_PER_INCH = 914400
CACHE = Path(os.environ.get("XDG_CACHE_HOME",
                            Path.home() / ".cache")) / "srs-help"


# ---------------------------------------------------------------------------
# Word field helpers — python-docx has no API for these
# ---------------------------------------------------------------------------
def add_field(paragraph, instr: str, initial: str = "1"):
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)

    r = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "
    r._r.append(it)

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._r.append(fld)

    paragraph.add_run(initial)

    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._r.append(fld)


# ---------------------------------------------------------------------------
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


class Renderer:
    def __init__(self, outline: dict, base: Path, mode: str = "merge",
                 draft: bool = False, root: Path = Path("."),
                 cfg: dict | None = None):
        self.o = outline
        self.cfg = cfg or {}
        self.mode = mode
        self.lex = outline["lexicon"]
        self.st = outline["styles"]
        self.base_lvl = outline["heading_base"][mode]
        self.draft = draft
        self.root = root
        self.doc = self._blank(base)
        self.warnings: list[str] = []
        self._missing_styles: set[str] = set()
        self.is_group = False
        self._puml_jar: Path | None = None

    @staticmethod
    def _blank(base: Path) -> Document:
        doc = Document(str(base))
        body = doc.element.body
        for ch in list(body):
            if ch.tag != qn("w:sectPr"):
                body.remove(ch)
        return doc

    def bullet_style(self, level: int) -> str:
        """Style for a bullet depth, falling back if the template lacks it.

        Level 3 needs a style the base template may not carry yet. Crashing on
        a missing style would block the whole document over a formatting
        detail, so fall back one level and say so once.
        """
        name = self.st.get(f"bullet_{min(level, 3)}") or self.st["bullet_1"]
        try:
            self.doc.styles[name]
            return name
        except KeyError:
            if name not in self._missing_styles:
                self._missing_styles.add(name)
                self.warnings.append(
                    f"mẫu Word chưa có style `{name}` cho gạch đầu dòng cấp "
                    f"{level} — tạm dùng cấp {level - 1}. Bổ sung style này "
                    f"vào cả mẫu con lẫn tài liệu tổng rồi render lại.")
            return self.bullet_style(level - 1) if level > 1 \
                else self.st["body"]

    # -- primitives --------------------------------------------------------
    def para(self, text: str = "", style: str | None = None, bold=False):
        p = self.doc.add_paragraph(style=style or self.st["body"])
        # Markdown emphasis has to become real runs. Passing the text through
        # verbatim puts literal asterisks in the document, which readers see as
        # a typo rather than emphasis.
        for part in INLINE_RE.split(text) if text else [""]:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            elif part.startswith("`") and part.endswith("`"):
                # No monospace style is allowed by the contract, so drop the
                # backticks and keep the text plain.
                p.add_run(part[1:-1])
            else:
                p.add_run(part)
        if bold:
            for r in p.runs:
                r.bold = True
        return p

    def heading(self, text: str, offset: int):
        lvl = max(1, min(self.base_lvl + offset, 9))
        return self.doc.add_paragraph(text, style=f"Heading {lvl}")

    def caption(self, kind: str, text: str):
        """kind: 'figure' | 'table'."""
        label = self.lex[kind]
        p = self.doc.add_paragraph(style=self.st["caption"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label} ")
        add_field(p, f"SEQ {label} \\* ARABIC")
        p.add_run(f". {text}" if text else "")

    # -- tables ------------------------------------------------------------
    def table(self, rows: list[list[str]], widths: list[int] | None,
              label: str = ""):
        if not rows:
            return
        # Render is the authority on STT: whatever the .md carries, the number
        # that reaches the document is the row's actual position. That way a
        # stale or missing ordinal in the source can never ship, and inserting
        # a row upstream needs no renumbering pass.
        rows = S.renumber_stt(rows)
        ncol = len(rows[0])
        widths = self._fit(widths, ncol)

        if label:
            self.para(label, bold=True)

        t = self.doc.add_table(rows=len(rows), cols=ncol)
        t.style = self.doc.styles[self.st["table"]]
        t.autofit = False
        head = [S.norm(c) for c in rows[0]] if rows else []
        allow = set(S.multiline_spec(self.o).get("cho_phep", []))
        for i, row in enumerate(rows):
            for j, val in enumerate(row[:ncol]):
                cell = t.rows[i].cells[j]
                col = head[j] if j < len(head) else ""
                if i > 0 and col in allow and S.cell_is_multiline(val):
                    self._cell_bullets(cell, val)
                else:
                    cell.text = val
                if i == 0:
                    for r in cell.paragraphs[0].runs:
                        r.bold = True
        for col, w in zip(t.columns, widths):
            col.width = Twips(w)
        for row in t.rows:
            self._cant_split(row)
            for cell, w in zip(row.cells, widths):
                cell.width = Twips(w)
        self.para("")
        return t

    def _cell_bullets(self, cell, val: str) -> None:
        """Write a `·`-marked cell as real bullet paragraphs.

        Only for columns the outline lists under `multiline_columns.cho_phep`
        — the sequential ones (Xử lý, Phản hồi của hệ thống, Kết quả / Mã
        thông báo), where several points are the nature of the content. The
        constraint columns stay on one line on purpose: the standard says each
        rule gets its own `BR-` code, and making a crowded cell read nicely
        would remove the reason to split it.

        Import reverses this exactly, so `.md → .docx → .md` is unchanged.
        """
        segs = S.cell_segments(val)
        first = True
        for text, level in segs:
            style = (self.st.get(f"bullet_{min(level, 3)}") if level
                     else self.st["body"])
            try:
                sobj = self.doc.styles[style]
            except KeyError:
                sobj = self.doc.styles[self.st["body"]]
            if first:
                cell.text = text
                cell.paragraphs[0].style = sobj
                first = False
            else:
                cell.add_paragraph(text, style=sobj)

    def _fit(self, widths: list[int] | None, ncol: int) -> list[int]:
        """Redistribute to the usable width when the column count differs.

        The permission matrix gains or loses role columns per function, so the
        outline widths are a starting point rather than a fixed contract.
        """
        usable = self.o["layout"]["usable_twips"]
        if widths and len(widths) == ncol:
            if sum(widths) == usable:
                return widths
            scale = usable / sum(widths)
            out = [int(w * scale) for w in widths]
        else:
            out = [usable // ncol] * ncol
        out[-1] += usable - sum(out)
        return out

    @staticmethod
    def _cant_split(row):
        trPr = row._tr.get_or_add_trPr()
        el = OxmlElement("w:cantSplit")
        trPr.append(el)

    # -- images ------------------------------------------------------------
    def image(self, path: Path, caption: str, alt_name: str = "",
              is_diagram: bool = False):
        if not path.exists():
            self.missing_box(caption or path.name,
                             f"không tìm thấy tệp `{path}`")
            return
        usable_emu = int(self.o["layout"]["usable_twips"] / 1440 * EMU_PER_INCH)
        try:
            from PIL import Image
            with Image.open(path) as im:
                px_w, px_h = im.size
        except Exception:
            px_w = px_h = 0

        min_w = self.o["images"]["min_width_px"]
        # The threshold guards against upscaling a low-resolution screenshot.
        # A PlantUML render is placed at its natural size and never stretched,
        # so applying it there would be a false alarm on every diagram.
        if px_w and px_w < min_w and not is_diagram:
            self.warnings.append(
                f"ảnh `{path.name}` rộng {px_w}px, dưới ngưỡng {min_w}px — "
                f"bản in sẽ mờ.")

        p = self.doc.add_paragraph(style=self.st["body"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if px_w and px_h:
            # Preserve the real aspect ratio: pass width only and let
            # python-docx derive the height. Forcing both distorts portrait
            # mockups, which is what the previous generator got wrong.
            width = min(usable_emu, int(px_w / 96 * EMU_PER_INCH))
            run.add_picture(str(path), width=Emu(width))
        else:
            run.add_picture(str(path), width=Emu(usable_emu))

        self._set_alt(run, alt_name or path.name)
        self.caption("figure", caption)

    @staticmethod
    def _set_alt(run, name: str):
        """Store the source filename in the picture's alt-text.

        Word does not keep the original filename, so import.py would otherwise
        have no way to restore assets/ with the right names on a round-trip.
        """
        for docPr in run._r.iter(qn("wp:docPr")):
            docPr.set("descr", name)
            docPr.set("name", name)

    def missing_box(self, what: str, why: str = ""):
        rows = [[f"{self.lex['missing_image']}  {what}"]]
        if why:
            rows.append([why])
        rows.append([
            "Bổ sung bằng 1 trong 2 cách: (1) gửi ảnh cho skill để thêm vào "
            "assets/; (2) chèn trực tiếp vào file .docx này, RỒI CHẠY IMPORT "
            "ngay sau đó — không import thì lần render kế tiếp sẽ mất ảnh."])
        self.table(rows, None)

    # -- PlantUML ----------------------------------------------------------
    def diagram(self, code: str, caption: str = ""):
        src = self.root / "diagrams" / f"{code}.puml"
        if not src.exists():
            self.missing_box(f"Sơ đồ {code}",
                             f"không tìm thấy `diagrams/{code}.puml`")
            return
        png = self._render_puml(src)
        if png is None:
            self.missing_box(f"Sơ đồ {code}",
                             "không render được PlantUML (thiếu Java hoặc mạng)")
            return
        self.image(png, caption or code, alt_name=f"{code}.png",
                   is_diagram=True)

    def _server_url(self) -> str | None:
        """Server address is environment-specific, so it lives in the project
        config next to the logo — not in the outline, which is the shared
        standard every project loads unchanged."""
        return (self.cfg.get("plantuml_server")
                or self.o["plantuml"].get("server"))

    def _render_puml(self, src: Path) -> Path | None:
        """Try the configured server first, then the local jar.

        A self-hosted server is only reachable from Claude Code — claude.ai
        restricts outbound hosts — so the jar remains the dependable path and
        the server is an opt-in shortcut for teams that run one.
        """
        cfg = self.o["plantuml"]
        outdir = CACHE / "png"
        outdir.mkdir(parents=True, exist_ok=True)
        dst = outdir / f"{src.stem}.png"

        server = self._server_url()
        if server:
            if self._via_server(server, src, dst):
                return dst
            self.warnings.append(
                f"không gọi được PlantUML server `{server}` — chuyển sang bản "
                f"jar cục bộ. Trên claude.ai server nội bộ luôn bị chặn; đây là "
                f"đường dùng cho Claude Code.")

        # Check Java before fetching 26 MB: without it the jar is useless, and
        # lumping the two causes into one message leaves the analyst guessing.
        if not shutil.which("java"):
            self._once("java",
                       "KHÔNG TÌM THẤY JAVA. PlantUML cần Java để chạy. Cài "
                       "JRE 17 trở lên (Windows: `winget install "
                       "Microsoft.OpenJDK.17`), hoặc khai `plantuml_server` "
                       "trong srs-config.json để render qua server thay vì máy "
                       "cục bộ.")
            return None

        jar = self._ensure_jar()
        if jar is None:
            return None
        try:
            subprocess.run(
                # -Sdpi raises the raster resolution; -scale is not a CLI
                # flag and is silently ignored, leaving diagrams too small
                # to read on a printed page.
                ["java", "-jar", str(jar), "-tpng", "-Sdpi=200",
                 "-o", str(outdir), str(src)],
                check=True, capture_output=True, timeout=180)
        except Exception as e:
            self.warnings.append(f"PlantUML lỗi với `{src.name}`: {e}")
            return None
        return dst if dst.exists() else None

    def _once(self, key: str, msg: str) -> None:
        """Report an environment problem once, not per diagram."""
        if key not in self._missing_styles:
            self._missing_styles.add(key)
            self.warnings.append(msg)

    @staticmethod
    def _encode_puml(text: str) -> str:
        """PlantUML's own deflate + base64 variant — not standard base64."""
        import zlib
        data = zlib.compress(text.encode("utf-8"))[2:-4]
        abc = ("0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"
               "abcdefghijklmnopqrstuvwxyz-_")
        out = []
        for i in range(0, len(data), 3):
            chunk = data[i:i + 3] + b"\x00" * (3 - len(data[i:i + 3]))
            n = (chunk[0] << 16) | (chunk[1] << 8) | chunk[2]
            out += [abc[(n >> 18) & 63], abc[(n >> 12) & 63],
                    abc[(n >> 6) & 63], abc[n & 63]]
        return "".join(out)

    def _via_server(self, server: str, src: Path, dst: Path) -> bool:
        try:
            enc = self._encode_puml(src.read_text(encoding="utf-8"))
            url = f"{server.rstrip('/')}/png/{enc}"
            with urllib.request.urlopen(url, timeout=30) as r:
                blob = r.read()
            if not blob.startswith(b"\x89PNG"):
                return False
            dst.write_bytes(blob)
            return True
        except Exception:
            return False

    def _ensure_jar(self) -> Path | None:
        if self._puml_jar is not None:
            return self._puml_jar
        cfg = self.o["plantuml"]
        ver = cfg["version"]

        # A jar the team placed by hand wins: it works offline and behind a
        # firewall that blocks github.com, which is the common case on a
        # corporate network.
        for cand in (self.cfg.get("plantuml_jar"),
                     "plantuml.jar", f"plantuml-{ver}.jar"):
            if not cand:
                continue
            p = Path(cand)
            if not p.is_absolute():
                p = self.root / cand
            if p.exists():
                self._puml_jar = p
                return p

        jar = CACHE / f"plantuml-{ver}.jar"
        if jar.exists():
            self._puml_jar = jar
            return jar
        url = (f"https://github.com/plantuml/plantuml/releases/download/"
               f"v{ver}/plantuml-{ver}.jar")
        try:
            CACHE.mkdir(parents=True, exist_ok=True)
            urllib.request.urlretrieve(url, jar)
            self._puml_jar = jar
            return jar
        except Exception as e:
            self._once("jar",
                       f"KHÔNG TẢI ĐƯỢC plantuml.jar ({type(e).__name__}: {e}). "
                       f"Mạng của bạn có thể chặn github.com. Tải tay từ\n"
                       f"    {url}\n"
                       f"  rồi đặt vào gốc dự án với tên `plantuml.jar`, hoặc "
                       f"khai đường dẫn ở `plantuml_jar` trong srs-config.json.")
            return None

    # -- cover page --------------------------------------------------------
    def cover(self, doc: S.FunctionDoc, cfg: dict):
        """Only for standalone output.

        In merge mode the child file is a fragment: the master document owns
        the cover, logo and running header. A fragment that carried its own
        would drop a logo into the middle of the merged document.
        """
        from docx.enum.text import WD_BREAK

        logo = cfg.get("logo")
        if logo:
            lp = (self.root / logo) if not Path(logo).is_absolute() else Path(logo)
            if lp.exists():
                p = self.doc.add_paragraph(style=self.st["body"])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                w = float(cfg.get("logo_rong_inch", 1.8))
                p.add_run().add_picture(str(lp), width=Emu(int(w * EMU_PER_INCH)))
            else:
                self.warnings.append(f"không tìm thấy logo `{lp}` — bỏ qua.")

        def line(text, size, bold=False, space=0):
            if not text:
                return
            p = self.doc.add_paragraph(style=self.st["body"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(size)
            if space:
                p.paragraph_format.space_after = Pt(space)

        line(cfg.get("to_chuc"), 13, True, 24)
        line(cfg.get("du_an"), 15, True, 6)
        line(cfg.get("tai_lieu", "Đặc tả yêu cầu phần mềm"), 20, True, 36)
        line(f"[{doc.ma}] {doc.meta.get('ten','')}", 16, True, 30)

        rows = [("Mã tài liệu", cfg.get("ma_tai_lieu", "")),
                ("Phiên bản", str(doc.meta.get("version", ""))),
                ("Ngày", self._last_date(doc)),
                ("Trạng thái", str(doc.meta.get("status", ""))),
                ("Bảo mật", cfg.get("bao_mat", ""))]
        rows = [r for r in rows if r[1]]
        if rows:
            usable = self.o["layout"]["usable_twips"]
            t = self.doc.add_table(rows=len(rows), cols=2)
            t.style = self.doc.styles[self.st["table"]]
            t.autofit = False
            for i, (k, v) in enumerate(rows):
                t.rows[i].cells[0].text = k
                t.rows[i].cells[1].text = v
                for r_ in t.rows[i].cells[0].paragraphs[0].runs:
                    r_.bold = True
            for col, w in zip(t.columns, [2400, usable - 2400]):
                col.width = Twips(w)
            for row in t.rows:
                for cell, w in zip(row.cells, [2400, usable - 2400]):
                    cell.width = Twips(w)

        p = self.doc.add_paragraph(style=self.st["body"])
        p.add_run().add_break(WD_BREAK.PAGE)

    @staticmethod
    def _last_date(doc: S.FunctionDoc) -> str:
        cl = doc.meta.get("changelog") or []
        for e in reversed(cl):
            if isinstance(e, dict) and e.get("ngay"):
                return str(e["ngay"])
        return ""

    def running_headers(self, cfg: dict):
        sec = self.doc.sections[0]
        txt = cfg.get("chan_trang") or cfg.get("du_an")
        if txt:
            hp = sec.header.paragraphs[0]
            hp.text = ""
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = hp.add_run(txt)
            r.font.size = Pt(10)
        if cfg.get("so_trang", True):
            fp = sec.footer.paragraphs[0]
            fp.text = ""
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_field(fp, "PAGE \\* ARABIC")

    # -- watermark ---------------------------------------------------------
    def stamp_draft(self):
        sec = self.doc.sections[0]
        hp = sec.header.paragraphs[0]
        if hp.text.strip():
            hp.add_run("  ·  ")
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run("BẢN NHÁP — CHƯA ĐỦ ĐIỀU KIỆN PHÁT HÀNH")
        r.bold = True
        r.font.size = Pt(11)

    # -- document ----------------------------------------------------------
    def render(self, doc: S.FunctionDoc) -> Document:
        prof = S.profile_of(self.o, doc.profile)
        self.is_group = bool(prof.get("is_group"))

        if self.mode == "standalone":
            self.cover(doc, self.cfg)
            self.running_headers(self.cfg)

        ten = doc.meta.get("ten", "")
        if self.is_group:
            # A group is a menu tier, not a document: one Heading 2 and an
            # optional short description, then straight into the functions.
            self.heading(f"Nhóm chức năng [{doc.ma}] {ten}", -1)
            for sec in doc.sections:
                for b in sec.blocks:
                    if b.kind == "note":
                        continue
                    if b.kind == "bullet":
                        self.para(b.text, style=self.bullet_style(b.level))
                    elif b.kind == "para":
                        # The approved master document styles a group's
                        # description as T-GhiChu, not body text.
                        self.para(b.text, style=self.st["note"])
            if self.draft:
                self.stamp_draft()
            return self.doc

        self.heading(f"Chức năng [{doc.ma}] {ten}", 0)

        feats_after = self._anchor(prof)
        for spec in prof["function_sections"]:
            sec = doc.section(spec["name"])
            self._section(spec, sec, doc, offset=1)
            if prof["has_features"] and spec["name"] == feats_after:
                for f in doc.features:
                    self.heading(f"Tính năng [{f.ma}] {f.ten}", 1)
                    for fspec in prof["feature_sections"]:
                        fs = next((x for x in f.sections
                                   if S.norm(x.name) == S.norm(fspec["name"])),
                                  None)
                        self._section(fspec, fs, doc, offset=2, feat=f)
        if self.draft:
            self.stamp_draft()
        return self.doc

    @staticmethod
    def _anchor(prof: dict) -> str:
        tail = {"Dữ liệu và tích hợp", "Phân loại dữ liệu",
                "Vấn đề còn mở", "Lịch sử thay đổi"}
        names = [s["name"] for s in prof["function_sections"]]
        for i in range(len(names) - 1, -1, -1):
            if names[i] not in tail:
                return names[i]
        return names[-1]

    def _section(self, spec: dict, sec: S.Section | None, doc: S.FunctionDoc,
                 offset: int, feat=None):
        if getattr(self, "is_group", False):
            # Bold body, not a heading: Heading 3 already belongs to the
            # function titles nested under this group.
            self.para(spec["name"], bold=True)
        else:
            self.heading(spec["name"], offset)

        if spec.get("auto") == "changelog":
            self._changelog(doc, spec)
            return
        if sec is None:
            self.para(self.lex["not_applicable"])
            return

        tspecs = [self.o["tables"][t] for t in spec.get("tables", [])]
        ti = 0
        for b in sec.blocks:
            if b.kind == "note":
                continue                      # guidance is for the BA, not the doc
            if b.kind == "tlabel":
                continue                      # emitted with its table
            if b.kind == "para":
                self.para(b.text)
            elif b.kind == "bullet":
                self.para(b.text, style=self.bullet_style(b.level))
            elif b.kind == "image":
                self.image(self.root / b.path, b.label)
            elif b.kind == "diagram":
                self.diagram(b.code)
            elif b.kind == "table":
                ts = tspecs[ti] if ti < len(tspecs) else None
                self.table(b.rows, ts["widths"] if ts else None,
                           label=b.label or (ts.get("label", "") if ts else ""))
                ti += 1

        if spec.get("visual") and not any(
                b.kind in ("image", "diagram") for b in sec.blocks):
            body = sec.text_content().lower()
            if not body.startswith(self.lex["not_applicable"].lower()):
                self.missing_box(spec["name"], "chưa có hình")

    def _changelog(self, doc: S.FunctionDoc, spec: dict):
        rows = [self.o["tables"][spec["tables"][0]]["headers"]]
        for e in doc.meta.get("changelog") or []:
            if isinstance(e, dict):
                rows.append([str(e.get("v", "")), str(e.get("ngay", "")),
                             str(e.get("nguoi", "")), str(e.get("mo_ta", ""))])
        self.table(rows, self.o["tables"][spec["tables"][0]]["widths"])


# ---------------------------------------------------------------------------
def gate_state(doc: S.FunctionDoc, outline: dict) -> list[str]:
    """Reasons this document may not be released, in the BA's words."""
    lex = outline["lexicon"]
    out = []

    text = []
    for sec in doc.sections:
        text.append(sec.text_content())
    for f in doc.features:
        for sec in f.sections:
            text.append(sec.text_content())
    n = "\n".join(text).count(lex["open_marker"])
    if n:
        out.append(f"còn {n} chỗ đánh dấu {lex['open_marker']} chưa chốt")

    vd = doc.section(outline["gate"]["section"])
    if vd:
        pend = sum(1 for b in vd.blocks if b.kind == "table"
                   for row in b.rows[1:]
                   if lex["status_pending"].lower() in " ".join(row).lower())
        if pend:
            out.append(f"mục “{outline['gate']['section']}” còn {pend} dòng "
                       f"“{lex['status_pending']}”")
    return out


def sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()[:16]


def main() -> int:
    S.utf8_stdio()
    ap = argparse.ArgumentParser(description="Sinh .docx từ file đặc tả .md.")
    ap.add_argument("src")
    ap.add_argument("-o", "--out", default=None)
    ap.add_argument("--standalone", action="store_true",
                    help="tài liệu độc lập (Heading 1) thay vì file con ghép "
                         "vào tài liệu tổng (Heading 3)")
    ap.add_argument("--draft", action="store_true",
                    help="ép đóng dấu BẢN NHÁP")
    ap.add_argument("--force-release", action="store_true",
                    help="bỏ qua cổng chặn và xuất bản phát hành dù còn điểm "
                         "treo — dùng khi cần gấp, tự chịu trách nhiệm")
    ap.add_argument("--outline", default=None)
    ap.add_argument("--base", default=None)
    ap.add_argument("--config", default=None,
                    help="cấu hình dự án; mặc định tìm ngược lên từ file .md")
    a = ap.parse_args()

    outline = S.load_outline(a.outline)
    src = Path(a.src)
    fdoc = S.read_markdown(src)
    if not S.known_kind(outline, fdoc.profile):
        print(f"LỖI: loại '{fdoc.profile}' không hợp lệ.",
              file=sys.stderr)
        return 1

    if a.out:
        out = Path(a.out)
    else:
        # merge.py finds child files by globbing `{MÃ}_*.docx`, so the name has
        # to carry the title too. Defaulting to just the code produces a file
        # the master document silently skips.
        ten = S.slug(fdoc.meta.get("ten", ""))
        ten = "-".join(w.capitalize() if i == 0 else w
                       for i, w in enumerate(ten.split("-")))
        out = src.with_name(f"{fdoc.ma}_{ten}.docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    # The gate has to be enforced here, not only in validate.py. Anyone who
    # skips validation would otherwise get a document that looks released while
    # still carrying unresolved marks.
    gate = gate_state(fdoc, outline)
    draft = a.draft or (bool(gate) and not a.force_release)
    if gate and not a.draft:
        for g in gate:
            print(f"  CỔNG CHẶN: {g}", file=sys.stderr)
        if a.force_release:
            print("  --force-release: vẫn xuất bản phát hành.", file=sys.stderr)
        else:
            print("  → Xuất BẢN NHÁP. Sạch hết rồi render lại để có bản phát "
                  "hành.", file=sys.stderr)

    cfg, cfg_path = S.find_config(src, a.config)
    if a.standalone and not cfg:
        print("  CẢNH BÁO: không tìm thấy srs-config.json — tài liệu độc lập "
              "sẽ không có bìa, logo, tên dự án. Chép "
              "assets/config.example.json ra gốc dự án.", file=sys.stderr)

    r = Renderer(outline, Path(a.base or S.BASE_DOCX),
                 mode="standalone" if a.standalone else "merge",
                 draft=draft, root=src.parent, cfg=cfg)
    d = r.render(fdoc)

    texts = [p_.text for p_ in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            texts += [c.text for c in row.cells]

    cp = d.core_properties
    cp.title = f"{fdoc.ma} {fdoc.meta.get('ten','')}"
    cp.comments = (f"srs-help | outline={outline['id']} v{outline['version']} "
                   f"| md_sha={sha(src)} | doc_sha={S.content_sha(texts)} "
                   f"| version={fdoc.meta.get('version')}")
    d.save(str(out))

    print(f"OK -> {out}"
          + ("   [BẢN NHÁP]" if draft else ""))
    want = "groups" if fdoc.profile == S.GROUP else f"functions/«phân hệ»"
    if out.parent.name not in ("groups",) and fdoc.profile == S.GROUP:
        print(f"  Ghi chú: để ghép vào tài liệu tổng, đặt file ở `{want}/`.")
    if a.standalone and cfg_path:
        print(f"  cấu hình: {cfg_path}")
    for w in r.warnings:
        print(f"  CẢNH BÁO: {w}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
